"""
Unit tests for role-aware color assignment in the formatting service.
"""
import os
import tempfile

import pytest
from docx import Document

from app.services.formatting_service import get_formatting_service


def _build_doc(path):
    doc = Document()
    doc.add_paragraph("My Main Title", style="Title")
    doc.add_paragraph("Chapter One", style="Heading 1")
    doc.add_paragraph("Section 1.1", style="Heading 2")
    doc.add_paragraph("Sub 1.1.1", style="Heading 3")
    doc.add_paragraph("Body text.")
    doc.save(path)


def _colors_by_style(path):
    doc = Document(path)
    result = {}
    for p in doc.paragraphs:
        rgb = None
        if p.runs and p.runs[0].font.color and p.runs[0].font.color.rgb is not None:
            rgb = str(p.runs[0].font.color.rgb)
        result.setdefault(p.style.name, rgb)
    return result


class TestFormattingColors:
    """Each enabled role must receive its own distinct color."""

    def test_four_roles_get_distinct_colors(self):
        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            out = os.path.join(d, "out.docx")
            _build_doc(inp)

            opts = {
                "titles": True,
                "section_titles": True,
                "paragraphs_titles": True,
                "paragraphs": True,
                "theme": {"positive": "#FF0000", "negative": "#0000FF", "scheme": "even"},
            }
            res = get_formatting_service().apply_formatting(inp, out, opts)
            assert res["success"] is True

            colors = _colors_by_style(out)
            title = colors["Title"]
            h1 = colors["Heading 1"]
            h2 = colors["Heading 2"]
            body = colors["Normal"]

            # None left uncolored (the previous 2-color model failed here).
            assert all(c is not None for c in (title, h1, h2, body))
            # All four roles distinct.
            assert len({title, h1, h2, body}) == 4

    def test_seed_colors_are_honored(self):
        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            out = os.path.join(d, "out.docx")
            _build_doc(inp)

            opts = {
                "titles": True,
                "paragraphs": True,
                "theme": {"positive": "#FF0000", "negative": "#0000FF"},
            }
            get_formatting_service().apply_formatting(inp, out, opts)
            colors = _colors_by_style(out)
            # titles -> first seed, paragraphs -> second seed (canonical order).
            assert colors["Title"] == "FF0000"
            assert colors["Normal"] == "0000FF"
