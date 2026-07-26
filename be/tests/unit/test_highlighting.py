"""
Unit tests for POS-based highlighting, focused on the two regressions
reported after a manual test session:

1. When highlighting finds nothing to change (no sections / no matching
   words), the output document must still be saved. Skipping the save left
   `formatted_path` pointing at a file that was never written, which broke
   every later operation chained from it (the rich preview fetch and the
   `/download` endpoint both failed, and `/download` returned a JSON error
   body that the browser saved as a bogus "file").
2. Highlighting must still reach paragraphs that were previously wrapped in
   a single-cell "framing" table by `apply_framing` (they are no longer part
   of `doc.paragraphs`, but are still real body content).
"""
import os
import tempfile

import pytest
from docx import Document
from docx.oxml.ns import qn

import app.services.keyword_service as ks
from app.services.formatting_service import get_formatting_service


class _FakeKeywordService:
    """Minimal POS tagger: every alphabetic word is tagged as a NOUN."""

    def analyze_pos(self, text):
        tokens = []
        idx = 0
        for word in text.split(' '):
            start = text.index(word, idx)
            end = start + len(word)
            is_punct = bool(word) and not any(c.isalnum() for c in word)
            tokens.append({
                'text': word,
                'pos': 'NOUN' if not is_punct else 'PUNCT',
                'lemma': word.lower(),
                'is_stop': False,
                'is_punct': is_punct,
                'is_space': not word.strip(),
                'start_char': start,
                'end_char': end,
            })
            idx = end
        return tokens

    def split_sentences(self, text):
        return [s for s in text.split('.') if s.strip()]


@pytest.fixture
def stub_keyword_service(monkeypatch):
    monkeypatch.setattr(ks, "_keyword_service_instance", _FakeKeywordService())
    yield


def _highlight_opts(**overrides):
    opts = {
        'enabled': True,
        'color': '#FF0000',
        'nouns': True,
        'verbs': False,
        'adjectives': False,
        'adverbs': False,
    }
    opts.update(overrides)
    return opts


class TestHighlightingSavesOutput:
    """Regression test: the document must always be written, even as a no-op."""

    def test_saves_document_even_when_no_pos_selected(self, stub_keyword_service):
        # A normal document with a heading and a body paragraph. The fake
        # tagger only ever emits NOUN tokens, so asking to highlight
        # adjectives (a valid, enabled request) matches nothing: nothing
        # should be highlighted, yet the output file must still exist and be
        # a valid, openable DOCX (formatted_path must never point at a
        # missing file).
        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            out = os.path.join(d, "out.docx")
            doc = Document()
            doc.add_paragraph("Chapter One", style="Heading 1")
            doc.add_paragraph("A simple body paragraph.")
            doc.save(inp)

            svc = get_formatting_service()
            result = svc.apply_highlighting(
                inp, out,
                _highlight_opts(nouns=False, adjectives=True),
            )

            assert result["success"] is True
            assert result["words_formatted"] == 0
            assert os.path.exists(out), (
                "Output file must exist even when nothing was highlighted"
            )
            # Must be a valid, openable DOCX.
            Document(out)

    def test_saves_document_when_no_sections_found(self, stub_keyword_service):
        # A document with no headings at all: section identification returns
        # nothing, but the fallback to top-level content paragraphs must
        # still kick in, highlight words, and save a valid file.
        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            out = os.path.join(d, "out.docx")
            doc = Document()
            doc.add_paragraph("Just a paragraph with no headings around it.")
            doc.save(inp)

            svc = get_formatting_service()
            result = svc.apply_highlighting(inp, out, _highlight_opts())

            assert result["success"] is True
            assert result["words_formatted"] > 0
            assert os.path.exists(out)
            Document(out)


class TestHighlightingReachesFramedContent:
    """Reproduces the reported scenario: colours -> framing -> highlighting."""

    def test_highlighting_after_framing_paragraphs(self, stub_keyword_service):
        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            framed = os.path.join(d, "framed.docx")
            highlighted = os.path.join(d, "highlighted.docx")

            doc = Document()
            doc.add_paragraph("Chapter One", style="Heading 1")
            doc.add_paragraph("First body paragraph of the section.")
            doc.add_paragraph("Second body paragraph of the section.")
            doc.save(inp)

            svc = get_formatting_service()

            # Step 1: frame the body paragraphs (default FE behaviour), which
            # moves them out of doc.paragraphs into 1x1 tables.
            frame_result = svc.apply_framing(
                inp, framed, {"paragraphs": True, "use_tables": True}
            )
            assert frame_result["success"] is True

            framed_doc = Document(framed)
            assert len(framed_doc.tables) == 2, (
                "Both body paragraphs should have been wrapped in tables"
            )

            # Step 2: highlighting must still find and colour the body text,
            # even though it now lives inside table cells.
            result = svc.apply_highlighting(framed, highlighted, _highlight_opts())

            assert result["success"] is True
            assert result["words_formatted"] > 0, (
                "Highlighting must reach paragraphs moved inside framing tables"
            )
            assert os.path.exists(highlighted)

            out_doc = Document(highlighted)
            # At least one run inside a table cell must carry the highlight color.
            highlighted_runs = [
                run
                for table in out_doc.tables
                for row in table.rows
                for cell in row.cells
                for para in cell.paragraphs
                for run in para.runs
                if run.font.color and run.font.color.rgb is not None
            ]
            assert highlighted_runs, "Expected at least one highlighted run inside a framed table cell"
            assert any(str(run.font.color.rgb) == "FF0000" for run in highlighted_runs)


