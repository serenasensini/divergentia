"""
Unit tests guarding against sensitive information (document text, AI prompts,
extracted keywords) leaking into application logs (issue #12).

These tests use a distinctive marker string that stands in for "sensitive
content" (e.g. personal information a user's document might contain) and
assert it never appears in any log record produced by the code paths that
handle document/text content, regardless of log level.
"""
import logging

import pytest
from unittest.mock import patch

from app.services.ollama_service import OllamaService
from app.services.keyword_service import get_keyword_service
from app.services.formatting_service import get_formatting_service

# A distinctive marker that would never legitimately appear in a log message
# unless the code was logging raw document/AI content.
SENSITIVE_MARKER = "Mario Rossi lives at Via Roma 42 and his SSN is 999-99-9999"


class TestOllamaServiceLogging:
    """Ollama prompts/responses embed document text and must never be logged."""

    def test_extract_keywords_does_not_log_prompt_or_response(self, app, caplog):
        with app.app_context():
            service = OllamaService()

            with patch.object(
                service, "_generate_completion", return_value=f"{SENSITIVE_MARKER}, other"
            ):
                with caplog.at_level(logging.DEBUG):
                    service.extract_keywords(
                        text=f"Some text about {SENSITIVE_MARKER}.",
                        max_keywords=5,
                        use_cache=False,
                    )

            for record in caplog.records:
                assert SENSITIVE_MARKER not in record.getMessage(), (
                    "Sensitive document content leaked into a log record: "
                    f"{record.getMessage()!r}"
                )

    def test_get_key_points_does_not_log_prompt_or_response(self, app, caplog):
        with app.app_context():
            service = OllamaService()

            with patch.object(
                service, "_generate_completion", return_value=f"{SENSITIVE_MARKER}\nAnother point"
            ):
                with caplog.at_level(logging.DEBUG):
                    service.get_key_points(text=f"Context about {SENSITIVE_MARKER}.", num_points=3)

            for record in caplog.records:
                assert SENSITIVE_MARKER not in record.getMessage(), (
                    "Sensitive document content leaked into a log record: "
                    f"{record.getMessage()!r}"
                )


class TestKeywordServiceLogging:
    """spaCy-based keyword extraction must not log raw text or keywords."""

    def test_extract_keywords_does_not_log_text(self, caplog):
        service = get_keyword_service()
        with caplog.at_level(logging.DEBUG):
            service.extract_keywords(f"Document about {SENSITIVE_MARKER} and other topics.")

        for record in caplog.records:
            assert SENSITIVE_MARKER not in record.getMessage()


class TestFormattingServiceLogging:
    """DOCX keyword-extraction pipeline must not log section/paragraph text."""

    def test_docx_keyword_extraction_does_not_log_document_text(self, caplog, tmp_path):
        from docx import Document

        input_path = tmp_path / "in.docx"
        output_path = tmp_path / "out.docx"

        doc = Document()
        doc.add_paragraph("Chapter One", style="Heading 1")
        doc.add_paragraph(f"This paragraph mentions {SENSITIVE_MARKER} explicitly.")
        doc.save(str(input_path))

        service = get_formatting_service()
        with caplog.at_level(logging.DEBUG):
            service.apply_keywords(
                str(input_path), str(output_path),
                {"max_keywords": 3, "include_proper_nouns": True},
            )

        for record in caplog.records:
            assert SENSITIVE_MARKER not in record.getMessage(), (
                "Sensitive document content leaked into a log record: "
                f"{record.getMessage()!r}"
            )


