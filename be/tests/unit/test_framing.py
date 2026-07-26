"""
Unit tests for table-based framing (sections / paragraphs / sentences).
"""
import os
import re
import tempfile

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import app.services.keyword_service as ks
from app.services.formatting_service import get_formatting_service


class _FakeKeywordService:
    """Lightweight sentence splitter so tests don't need the spaCy model."""

    def split_sentences(self, text):
        return [s.strip() for s in re.split(r'(?<=[.!?])\s+', text.strip()) if s.strip()]


@pytest.fixture
def stub_sentences(monkeypatch):
    monkeypatch.setattr(ks, "_keyword_service_instance", _FakeKeywordService())
    yield


def _build_doc(path):
    doc = Document()
    doc.add_paragraph("Chapter One", style="Heading 1")
    doc.add_paragraph("First body paragraph of the section.")
    doc.add_paragraph("A standalone paragraph. It has two sentences here.")
    doc.save(path)


def _frame(opts):
    with tempfile.TemporaryDirectory() as d:
        inp = os.path.join(d, "in.docx")
        out = os.path.join(d, "out.docx")
        _build_doc(inp)
        res = get_formatting_service().apply_framing(inp, out, {**opts, "use_tables": True})
        assert res["success"] is True
        return Document(out)


def _body_elements(doc):
    return list(doc.element.body)


def _has_adjacent_tables(doc):
    body = _body_elements(doc)
    return any(
        body[i].tag == qn('w:tbl') and body[i + 1].tag == qn('w:tbl')
        for i in range(len(body) - 1)
    )


def _top_border(table):
    b = table._element.find(qn('w:tblPr')).find(qn('w:tblBorders')).find(qn('w:top'))
    return b.get(qn('w:val')), b.get(qn('w:sz'))


def _side_val(table, side):
    """Return the ``w:val`` of a single outer border (e.g. 'double' or 'nil')."""
    borders = table._element.find(qn('w:tblPr')).find(qn('w:tblBorders'))
    return borders.find(qn(f'w:{side}')).get(qn('w:val'))


class TestFraming:
    def test_sections_merge_into_single_table(self):
        doc = _frame({"sections": True})
        assert len(doc.tables) == 1
        cell = doc.tables[0].rows[0].cells[0]
        # Both body paragraphs of the section end up in the one cell.
        assert len(cell.paragraphs) == 2
        assert _top_border(doc.tables[0]) == ("double", "16")
        assert not _has_adjacent_tables(doc)

    def test_paragraphs_one_table_each(self):
        doc = _frame({"paragraphs": True})
        assert len(doc.tables) == 2
        for t in doc.tables:
            assert len(t.rows[0].cells[0].paragraphs) == 1
            assert _top_border(t) == ("single", "8")
        assert not _has_adjacent_tables(doc)

    def test_sentences_one_table_each(self, stub_sentences):
        doc = _frame({"sentences": True})
        # 1 sentence + 2 sentences = 3 tables.
        assert len(doc.tables) == 3
        texts = [t.rows[0].cells[0].text for t in doc.tables]
        assert texts == [
            "First body paragraph of the section.",
            "A standalone paragraph.",
            "It has two sentences here.",
        ]
        for t in doc.tables:
            assert _top_border(t) == ("dashed", "4")
        assert not _has_adjacent_tables(doc)

    def test_precedence_sections_over_paragraphs_and_sentences(self, stub_sentences):
        doc = _frame({"sections": True, "paragraphs": True, "sentences": True})
        # Every body paragraph belongs to the section, so only the section frame
        # is produced (no double-wrapping).
        assert len(doc.tables) == 1
        assert _top_border(doc.tables[0]) == ("double", "16")

    def test_custom_border_overrides_defaults(self):
        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            out = os.path.join(d, "out.docx")
            _build_doc(inp)
            get_formatting_service().apply_framing(
                inp, out,
                {"paragraphs": True, "use_tables": True,
                 "border_style": "dotted", "border_width": 12, "border_color": "#FF0000"},
            )
            doc = Document(out)
            val, sz = _top_border(doc.tables[0])
            assert val == "dotted"
            assert sz == "12"
            color = doc.tables[0]._element.find(qn('w:tblPr')).find(
                qn('w:tblBorders')).find(qn('w:top')).get(qn('w:color'))
            assert color == "FF0000"


class TestFramingAcrossPageBreak:
    """A framed block that spans a page break becomes an open-ended box."""

    def test_section_split_via_page_break_before(self):
        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            out = os.path.join(d, "out.docx")
            doc = Document()
            doc.add_paragraph("Chapter One", style="Heading 1")
            doc.add_paragraph("First part, before the page break.")
            cont = doc.add_paragraph("Second part, on the next page.")
            cont.paragraph_format.page_break_before = True
            doc.save(inp)

            res = get_formatting_service().apply_framing(
                inp, out, {"sections": True, "use_tables": True})
            assert res["success"] is True

            out_doc = Document(out)
            assert len(out_doc.tables) == 2
            first, second = out_doc.tables

            # First part: box open at the bottom (top + sides drawn).
            assert _side_val(first, 'top') == 'double'
            assert _side_val(first, 'left') == 'double'
            assert _side_val(first, 'right') == 'double'
            assert _side_val(first, 'bottom') == 'nil'

            # Continuation: box open at the top (bottom + sides drawn).
            assert _side_val(second, 'top') == 'nil'
            assert _side_val(second, 'bottom') == 'double'
            assert _side_val(second, 'left') == 'double'
            assert _side_val(second, 'right') == 'double'

    def test_paragraph_split_via_explicit_page_break(self):
        from docx.enum.text import WD_BREAK

        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            out = os.path.join(d, "out.docx")
            doc = Document()
            doc.add_paragraph("Title", style="Heading 1")
            para = doc.add_paragraph("Text before the break.")
            para.add_run().add_break(WD_BREAK.PAGE)
            para.add_run(" Text after the break.")
            doc.save(inp)

            res = get_formatting_service().apply_framing(
                inp, out, {"paragraphs": True, "use_tables": True})
            assert res["success"] is True

            out_doc = Document(out)
            assert len(out_doc.tables) == 2
            first, second = out_doc.tables
            assert _side_val(first, 'top') == 'single'
            assert _side_val(first, 'bottom') == 'nil'
            assert _side_val(second, 'top') == 'nil'
            assert _side_val(second, 'bottom') == 'single'

    def test_single_page_block_keeps_full_box(self):
        # No page break: behaviour is unchanged (one closed box).
        doc = _frame({"sections": True})
        assert len(doc.tables) == 1
        t = doc.tables[0]
        for side in ('top', 'bottom', 'left', 'right'):
            assert _side_val(t, side) == 'double'


class TestFramingLists:
    """A contiguous run of list items is framed as a single box (issue #14)."""

    @staticmethod
    def _add_num_pr(paragraph, num_id="1", ilvl="0"):
        """Attach direct ``w:numPr`` to a paragraph, as Word does for real lists."""
        pPr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement('w:numPr')
        ilvl_el = OxmlElement('w:ilvl')
        ilvl_el.set(qn('w:val'), ilvl)
        num_id_el = OxmlElement('w:numId')
        num_id_el.set(qn('w:val'), num_id)
        num_pr.append(ilvl_el)
        num_pr.append(num_id_el)
        pPr.append(num_pr)

    def _build_doc_with_list(self, path, list_style="List Bullet"):
        doc = Document()
        doc.add_paragraph("Chapter One", style="Heading 1")
        doc.add_paragraph("Intro paragraph before the list.")
        for text in ("First item", "Second item", "Third item"):
            p = doc.add_paragraph(text, style=list_style)
            self._add_num_pr(p)
        doc.add_paragraph("Outro paragraph after the list.")
        doc.save(path)

    def test_list_items_grouped_into_single_table(self):
        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            out = os.path.join(d, "out.docx")
            self._build_doc_with_list(inp)

            res = get_formatting_service().apply_framing(
                inp, out, {"paragraphs": True, "use_tables": True})
            assert res["success"] is True

            doc = Document(out)
            # 2 non-list paragraphs framed individually + 1 table for the
            # whole list (not one table per list item).
            assert len(doc.tables) == 3

            list_table = next(
                t for t in doc.tables if len(t.rows[0].cells[0].paragraphs) == 3
            )
            cell = list_table.rows[0].cells[0]
            assert [p.text for p in cell.paragraphs] == [
                "First item", "Second item", "Third item",
            ]

    def test_list_numbering_preserved_when_framed(self):
        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            out = os.path.join(d, "out.docx")
            self._build_doc_with_list(inp)

            res = get_formatting_service().apply_framing(
                inp, out, {"paragraphs": True, "use_tables": True})
            assert res["success"] is True

            doc = Document(out)
            list_table = next(
                t for t in doc.tables if len(t.rows[0].cells[0].paragraphs) == 3
            )
            for p in list_table.rows[0].cells[0].paragraphs:
                pPr = p._p.pPr
                assert pPr is not None
                assert pPr.find(qn('w:numPr')) is not None, (
                    "List numbering (w:numPr) should be preserved when the "
                    "list is encapsulated in a table"
                )

    def test_list_split_across_page_break_reuses_segmentation(self):
        with tempfile.TemporaryDirectory() as d:
            inp = os.path.join(d, "in.docx")
            out = os.path.join(d, "out.docx")
            doc = Document()
            doc.add_paragraph("Chapter One", style="Heading 1")
            p1 = doc.add_paragraph("First item", style="List Bullet")
            self._add_num_pr(p1)
            second = doc.add_paragraph("Second item", style="List Bullet")
            self._add_num_pr(second)
            second.paragraph_format.page_break_before = True
            p3 = doc.add_paragraph("Third item", style="List Bullet")
            self._add_num_pr(p3)
            doc.save(inp)

            res = get_formatting_service().apply_framing(
                inp, out, {"paragraphs": True, "use_tables": True})
            assert res["success"] is True

            out_doc = Document(out)
            # The list run spans a page break: it becomes 2 open-ended boxes
            # instead of one, exactly like sections/paragraphs already do.
            assert len(out_doc.tables) == 2
            first, second_table = out_doc.tables
            assert _side_val(first, 'bottom') == 'nil'
            assert _side_val(second_table, 'top') == 'nil'


