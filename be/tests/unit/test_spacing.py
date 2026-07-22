"""
Unit tests for sentence/paragraph spacing.

Regression coverage for the bug where sentence spacing split text word-by-word
and inserted a doubled blank-line gap between sentences.
"""
import re

import pytest
from docx import Document
from docx.oxml.ns import qn

import app.services.keyword_service as ks
from app.services.formatting_service import get_formatting_service


class _FakeKeywordService:
    """Lightweight sentence splitter so tests don't need the spaCy model."""

    def split_sentences(self, text):
        return [s.strip() for s in re.split(r'(?<=[.!?])\s+', text.strip()) if s.strip()]


@pytest.fixture
def stub_sentences(monkeypatch):
    monkeypatch.setattr(ks, "_keyword_service_instance", _FakeKeywordService())
    yield


def _run_sentence_spacing(text):
    doc = Document()
    p = doc.add_paragraph(text)
    get_formatting_service()._add_sentence_spacing(p)
    return p


def test_sentence_spacing_preserves_words(stub_sentences):
    text = "Questa è la prima frase. Ecco la seconda frase! E la terza?"
    p = _run_sentence_spacing(text)

    # No word-splitting: replacing the inserted breaks with spaces
    # reconstructs the original text exactly.
    assert p.text.replace("\n", " ") == text


def test_sentence_spacing_uses_single_break(stub_sentences):
    text = "First sentence here. Second sentence here. Third one."
    p = _run_sentence_spacing(text)

    xml = p._p.xml
    # One break per sentence gap (3 sentences -> 2 breaks), never doubled.
    assert xml.count("<w:br/>") == 2
    assert "\n\n" not in p.text
    assert "\t" not in p.text


def test_single_sentence_is_untouched(stub_sentences):
    text = "Just one sentence."
    p = _run_sentence_spacing(text)

    assert "<w:br/>" not in p._p.xml
    assert p.text == text


def test_spacing_response_is_compact(tmp_path, stub_sentences):
    inp = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("Prima frase qui. Seconda frase qui.")
    doc.save(inp)

    res = get_formatting_service()._apply_spacing_docx(
        str(inp), str(out), {"paragraphs": False, "sentences": True}
    )

    assert set(res.keys()) == {"success", "output_path", "format", "spacing_applied"}
    assert res["success"] is True


def test_list_item_gets_padding_not_line_breaks(stub_sentences):
    svc = get_formatting_service()
    doc = Document()
    li = doc.add_paragraph("Primo elemento. Con due frasi.", style="List Number")

    svc._add_sentence_spacing(li)

    # No line breaks inserted; text left intact.
    assert "<w:br/>" not in li._p.xml
    assert li.text == "Primo elemento. Con due frasi."
    # Padding above and below the element (~10px => 7.5pt).
    assert li.paragraph_format.space_before.pt == pytest.approx(7.5)
    assert li.paragraph_format.space_after.pt == pytest.approx(7.5)


def test_bulleted_list_item_paragraph_spacing_uses_padding(stub_sentences):
    svc = get_formatting_service()
    doc = Document()
    li = doc.add_paragraph("Voce di elenco puntato.", style="List Bullet")

    assert svc._is_list_paragraph(li) is True

    svc._add_paragraph_spacing(li)

    # Padding instead of surrounding <w:br/> siblings.
    assert li.paragraph_format.space_before.pt == pytest.approx(7.5)
    assert li.paragraph_format.space_after.pt == pytest.approx(7.5)
    p = li._element
    assert p.getprevious() is None or p.getprevious().tag != qn("w:br")
    assert p.getnext() is None or p.getnext().tag != qn("w:br")
    # Contextual spacing disabled so the padding renders between items.
    cs = li._p.pPr.find(qn("w:contextualSpacing"))
    assert cs is not None and cs.get(qn("w:val")) == "0"


def test_plain_paragraph_is_not_a_list(stub_sentences):
    svc = get_formatting_service()
    doc = Document()
    p = doc.add_paragraph("Un paragrafo normale.")
    assert svc._is_list_paragraph(p) is False
