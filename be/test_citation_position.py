"""
Test specifico per verificare che i run con citazioni siano preservati nell'ordine corretto
"""
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.services.formatting_service import FormattingService

# Create a test document with a simulated citation field
doc = Document()
para = doc.add_paragraph()

# Add "The noun "
run1 = para.add_run("The noun ")
run1.font.color.rgb = RGBColor(0, 0, 0)

# Add a run with fldChar to simulate a citation - THIS SHOULD BE SKIPPED AND STAY IN PLACE
run2 = para.add_run("[Citation]")
run2.font.color.rgb = RGBColor(128, 0, 128)  # Purple for citation
run2.italic = True

# Manually add a fldChar element to simulate a real citation field
fld_char = OxmlElement('w:fldChar')
fld_char.set(qn('w:fldCharType'), 'begin')
run2._element.append(fld_char)

# Add " is a noun"
run3 = para.add_run(" is a noun")
run3.font.color.rgb = RGBColor(0, 0, 0)

test_input_path = "/home/ssensini/WebstormProjects/divergentia/be/outputs/test_citation_field.docx"
doc.save(test_input_path)

print("="*80)
print("TEST: PRESERVAZIONE POSIZIONE CITAZIONI CON CAMPI COMPLESSI")
print("="*80)
print(f"Testo originale: {para.text}")
print(f"Numero di run: {len(para.runs)}")
for i, run in enumerate(para.runs):
    has_field = False
    if run._element is not None:
        for child in run._element:
            if 'fldChar' in str(child.tag):
                has_field = True
                break
    field_marker = " [HAS FLDCHAR]" if has_field else ""
    color = run.font.color.rgb if run.font.color.rgb else None
    print(f"  Run {i}: '{run.text}' - color: {color}, italic: {run.italic}{field_marker}")

print()
print("Applicazione highlighting per SOSTANTIVI...")
print("-"*80)

formatting_service = FormattingService()
output_path = "/home/ssensini/WebstormProjects/divergentia/be/outputs/test_citation_field_output.docx"

result = formatting_service.apply_highlighting(
    test_input_path,
    output_path,
    {
        'enabled': True,
        'color': '#FF0000',  # Red for nouns
        'style': 'bold',
        'font_size': None,
        'font_family': None,
        'nouns': True,
        'verbs': False,
        'adjectives': False,
        'adverbs': False
    }
)

print(f"✓ Highlighting applicato")
print(f"  Parole formattate: {result['words_formatted']}")
print(f"  Sostantivi: {result['pos_stats']['nouns']}")

print()
print("DOCUMENTO OUTPUT:")
print("-"*80)

output_doc = Document(output_path)
output_para = output_doc.paragraphs[0]

print(f"Testo: {output_para.text}")
print(f"Numero di run: {len(output_para.runs)}")

for i, run in enumerate(output_para.runs):
    has_field = False
    if run._element is not None:
        for child in run._element:
            if 'fldChar' in str(child.tag):
                has_field = True
                break

    field_marker = " [HAS FLDCHAR]" if has_field else ""
    color_str = "ROSSO" if run.font.color.rgb == RGBColor(255, 0, 0) else \
                "VIOLA" if run.font.color.rgb == RGBColor(128, 0, 128) else \
                "NERO" if run.font.color.rgb == RGBColor(0, 0, 0) else "ALTRO"
    bold_str = " bold" if run.bold else ""
    italic_str = " italic" if run.italic else ""

    print(f"  Run {i}: '{run.text}' - {color_str}{bold_str}{italic_str}{field_marker}")

print()
print("="*80)
print("VERIFICA")
print("="*80)

# Check if citation is in the middle (not at the beginning or end)
citation_position = None
for i, run in enumerate(output_para.runs):
    if '[Citation]' in run.text:
        citation_position = i
        break

if citation_position is None:
    print("✗ ERRORE: Citazione non trovata!")
elif citation_position == 0:
    print(f"✗ ERRORE: Citazione spostata all'inizio! (posizione: {citation_position})")
elif citation_position == len(output_para.runs) - 1:
    print(f"✗ ERRORE: Citazione spostata alla fine! (posizione: {citation_position})")
else:
    print(f"✓ Citazione rimasta nel mezzo del paragrafo (posizione: {citation_position}/{len(output_para.runs)})")

# Check if citation has field marker
citation_run = output_para.runs[citation_position] if citation_position is not None else None
if citation_run:
    has_field = False
    if citation_run._element is not None:
        for child in citation_run._element:
            if 'fldChar' in str(child.tag):
                has_field = True
                break

    if has_field:
        print("✓ Citazione ha ancora il campo fldChar")
    else:
        print("✗ Citazione ha perso il campo fldChar")

    # Check color
    if citation_run.font.color.rgb == RGBColor(128, 0, 128):
        print("✓ Citazione ha mantenuto il colore viola")
    else:
        print(f"✗ Citazione ha cambiato colore: {citation_run.font.color.rgb}")

print()
print("Atteso: La citazione dovrebbe rimanere tra 'The noun ' e ' is a noun'")
print("        e NON essere formattata (rimanere viola, italic, senza bold)")

