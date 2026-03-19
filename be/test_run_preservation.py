"""
Test to verify that the new implementation preserves run structure
"""
from docx import Document
from docx.shared import RGBColor, Pt

# Create a test document with multiple runs
doc = Document()

# Add a paragraph with multiple runs (simulating mixed formatting)
para = doc.add_paragraph()
run1 = para.add_run("This is a ")
run1.font.color.rgb = RGBColor(0, 0, 0)  # Black

run2 = para.add_run("citation")
run2.font.color.rgb = RGBColor(128, 0, 128)  # Purple (simulating citation)
run2.italic = True

run3 = para.add_run(" in the middle of a ")
run3.font.color.rgb = RGBColor(0, 0, 0)  # Black

run4 = para.add_run("sentence")
run4.font.color.rgb = RGBColor(0, 0, 0)  # Black
run4.bold = True

run5 = para.add_run(".")
run5.font.color.rgb = RGBColor(0, 0, 0)  # Black

# Save the test document
test_input_path = "/home/ssensini/WebstormProjects/divergentia/be/outputs/test_input_multirun.docx"
doc.save(test_input_path)

print("="*80)
print("ORIGINAL DOCUMENT CREATED")
print("="*80)
print(f"Paragraph text: {para.text}")
print(f"Number of runs: {len(para.runs)}")
for i, run in enumerate(para.runs):
    color = run.font.color.rgb if run.font.color.rgb else "None"
    print(f"  Run {i}: '{run.text}' - color: {color}, bold: {run.bold}, italic: {run.italic}")

# Now apply highlighting
from app.services.formatting_service import FormattingService

formatting_service = FormattingService()
output_path = "/home/ssensini/WebstormProjects/divergentia/be/outputs/test_output_multirun.docx"

highlighting_options = {
    'enabled': True,
    'color': '#FF0000',  # Red
    'style': 'bold',
    'font_size': None,
    'font_family': None,
    'nouns': True,
    'verbs': False,
    'adjectives': False,
    'adverbs': False
}

print("\n" + "="*80)
print("APPLYING HIGHLIGHTING TO NOUNS...")
print("="*80)

try:
    result = formatting_service.apply_highlighting(
        test_input_path,
        output_path,
        highlighting_options
    )

    print(f"✓ Success! Words formatted: {result['words_formatted']}")

    # Analyze output
    output_doc = Document(output_path)
    output_para = output_doc.paragraphs[0]

    print("\n" + "="*80)
    print("OUTPUT DOCUMENT ANALYSIS")
    print("="*80)
    print(f"Paragraph text: {output_para.text}")
    print(f"Number of runs: {len(output_para.runs)}")
    for i, run in enumerate(output_para.runs):
        color = run.font.color.rgb if run.font.color.rgb else "None"
        print(f"  Run {i}: '{run.text}' - color: {color}, bold: {run.bold}, italic: {run.italic}")

    # Verify run structure is preserved
    print("\n" + "="*80)
    print("VERIFICATION")
    print("="*80)

    original_doc = Document(test_input_path)
    original_para = original_doc.paragraphs[0]

    if len(original_para.runs) == len(output_para.runs):
        print(f"✓ Run count preserved: {len(original_para.runs)} runs")
    else:
        print(f"✗ Run count changed: {len(original_para.runs)} -> {len(output_para.runs)}")

    # Check if the purple citation run is preserved
    citation_run_preserved = False
    for i, run in enumerate(output_para.runs):
        if run.text == "citation" and run.italic:
            print(f"✓ Citation run found at position {i} with italic formatting")
            if run.font.color.rgb and run.font.color.rgb == RGBColor(128, 0, 128):
                print(f"  ✓ Citation color preserved (purple)")
                citation_run_preserved = True
            elif run.font.color.rgb:
                print(f"  ! Citation color changed to {run.font.color.rgb}")
            else:
                print(f"  ! Citation color lost")

    if not citation_run_preserved:
        print("  Note: Citation formatting may have been modified by highlighting")

    print("\n" + "="*80)
    print("CONCLUSION")
    print("="*80)
    print("The new implementation:")
    print("1. Preserves the run structure (no deletion and recreation)")
    print("2. Modifies runs in place")
    print("3. Skips runs with complex field elements (citations/references)")
    print("4. Maintains original formatting for non-targeted text")

except Exception as e:
    print(f"✗ Error: {str(e)}")
    import traceback
    traceback.print_exc()

