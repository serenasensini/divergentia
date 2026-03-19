"""
Simple test to verify citations are skipped
"""
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.services.formatting_service import FormattingService

print("="*80)
print("TEST: Paragraphs with citations should be SKIPPED")
print("="*80)

# Create test document
doc = Document()

# Paragraph 1: Normal text WITHOUT citations
para1 = doc.add_paragraph("This paragraph has nouns like cat and dog.")

# Paragraph 2: Text WITH a citation field
para2 = doc.add_paragraph()
run1 = para2.add_run("This has nouns ")
run_citation = para2.add_run("(Author, 2020)")
run_citation.font.color.rgb = RGBColor(128, 0, 128)
# Add fldChar to simulate citation
fld_char = OxmlElement('w:fldChar')
fld_char.set(qn('w:fldCharType'), 'begin')
run_citation._element.append(fld_char)
run2 = para2.add_run(" and more nouns here.")

test_input = '/home/ssensini/WebstormProjects/divergentia/be/outputs/test_skip_citations.docx'
test_output = '/home/ssensini/WebstormProjects/divergentia/be/outputs/test_skip_citations_output.docx'
doc.save(test_input)

print("\nInput document:")
print(f"  Para 0: '{para1.text}' - {len(para1.runs)} runs")
print(f"  Para 1: '{para2.text}' - {len(para2.runs)} runs (HAS CITATION)")

print("\nApplying formatting...")

service = FormattingService()
result = service.apply_highlighting(
    test_input,
    test_output,
    {
        'enabled': True,
        'color': '#FF0000',
        'style': 'bold',
        'nouns': True,
        'verbs': False,
        'adjectives': False,
        'adverbs': False
    }
)

print(f"\n✓ Done!")
print(f"  Words formatted: {result['words_formatted']}")
print(f"  Paragraphs processed: {result['paragraphs_processed']}")

# Check output
output_doc = Document(test_output)

print("\nOutput document:")
para1_out = output_doc.paragraphs[0]
para2_out = output_doc.paragraphs[1]

print(f"  Para 0: {len(para1_out.runs)} runs")
if len(para1_out.runs) > 1:
    print("    ✓ Paragraph WITHOUT citation WAS formatted")
else:
    print("    ✗ Paragraph WITHOUT citation was NOT formatted")

print(f"  Para 1: {len(para2_out.runs)} runs")
if len(para2_out.runs) == 3:  # Same as input
    print("    ✓ Paragraph WITH citation was SKIPPED (not formatted)")
    # Check citation still there
    has_citation = any('fldChar' in str(c.tag)
                      for run in para2_out.runs
                      for c in run._element if run._element)
    if has_citation:
        print("    ✓ Citation field preserved")
    else:
        print("    ✗ Citation field lost")
else:
    print(f"    ✗ Paragraph WITH citation was formatted (runs changed from 3 to {len(para2_out.runs)})")

print("\n" + "="*80)
print("EXPECTED: Para 0 formatted, Para 1 skipped (citation preserved)")
print("="*80)

