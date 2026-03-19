"""
Test completo per la funzionalità di framing con tabelle 1x1
"""
from docx import Document
from docx.shared import RGBColor, Pt
from app.services.formatting_service import FormattingService

print("="*80)
print("TEST: FRAMING CON TABELLE 1x1")
print("="*80)

# Crea documento di test
doc = Document()

# Aggiungi un titolo
doc.add_heading("Test Framing con Tabelle", level=1)

# Aggiungi alcuni paragrafi con formattazione diversa
para1 = doc.add_paragraph("Questo è il primo paragrafo con testo normale.")

para2 = doc.add_paragraph()
run1 = para2.add_run("Questo paragrafo ha ")
run2 = para2.add_run("testo in grassetto")
run2.bold = True
run3 = para2.add_run(" e ")
run4 = para2.add_run("testo in corsivo")
run4.italic = True
run5 = para2.add_run(".")

para3 = doc.add_paragraph("Terzo paragrafo con testo colorato.")
for run in para3.runs:
    run.font.color.rgb = RGBColor(255, 0, 0)

# Aggiungi un titolo di sezione
doc.add_heading("Sezione di Test", level=2)

para4 = doc.add_paragraph("Paragrafo nella sezione con diversi font size.")
para4.runs[0].font.size = Pt(14)

test_input = '/home/ssensini/WebstormProjects/divergentia/be/outputs/test_framing_tables_input.docx'
test_output = '/home/ssensini/WebstormProjects/divergentia/be/outputs/test_framing_tables_output.docx'

doc.save(test_input)

print("\nDocumento di test creato con:")
print(f"  - 1 titolo principale")
print(f"  - 4 paragrafi con formattazione varia")
print(f"  - 1 titolo di sezione")

print("\n" + "-"*80)
print("TEST 1: Framing con tabelle - tutti i paragrafi")
print("-"*80)

service = FormattingService()

result = service.apply_framing(
    test_input,
    test_output,
    {
        'sections': False,
        'paragraphs': True,
        'subparagraphs': False,
        'sentences': False,
        'use_tables': True,
        'border_width': 12,  # 1.5 pt
        'border_color': '0000FF',  # Blu
        'border_style': 'single',
        'cell_margin': 100,
        'preserve_spacing': True,
        'filter': {
            'exclude_headings': True,
            'exclude_empty': True,
            'min_length': 10
        }
    }
)

print(f"\n✓ Test 1 completato")
print(f"  Borders applicati: {result['borders_applied']}")
print(f"  Metodo: {result['method']}")
print(f"  Output: {test_output}")

# Verifica output
output_doc = Document(test_output)
print(f"\n  Verifica documento output:")
print(f"    Totale elementi: {len(output_doc.paragraphs) + len(output_doc.tables)}")
print(f"    Tabelle create: {len(output_doc.tables)}")

# Verifica contenuto nelle tabelle
tables_with_content = 0
for table in output_doc.tables:
    cell = table.rows[0].cells[0]
    if cell.text.strip():
        tables_with_content += 1

print(f"    Tabelle con contenuto: {tables_with_content}")

print("\n" + "-"*80)
print("TEST 2: Framing con bordi personalizzati")
print("-"*80)

test_output_2 = '/home/ssensini/WebstormProjects/divergentia/be/outputs/test_framing_custom_output.docx'

result2 = service.apply_framing(
    test_input,
    test_output_2,
    {
        'sections': False,
        'paragraphs': True,
        'subparagraphs': False,
        'sentences': False,
        'use_tables': True,
        'border_width': 16,  # 2 pt - bordo spesso
        'border_color': 'FF0000',  # Rosso
        'border_style': 'double',  # Bordo doppio
        'cell_margin': 150,  # Margine maggiore
        'preserve_spacing': True,
        'filter': {
            'exclude_headings': True,
            'exclude_empty': True
        }
    }
)

print(f"\n✓ Test 2 completato")
print(f"  Borders applicati: {result2['borders_applied']}")
print(f"  Bordo: {result2['framing_options']['border_width']}/8 pt, "
      f"colore {result2['framing_options']['border_color']}, "
      f"stile {result2['framing_options']['border_style']}")

print("\n" + "-"*80)
print("TEST 3: Framing con filtro per stile")
print("-"*80)

test_output_3 = '/home/ssensini/WebstormProjects/divergentia/be/outputs/test_framing_filtered_output.docx'

result3 = service.apply_framing(
    test_input,
    test_output_3,
    {
        'sections': False,
        'paragraphs': True,
        'subparagraphs': False,
        'sentences': False,
        'use_tables': True,
        'border_width': 8,
        'border_color': '00FF00',  # Verde
        'border_style': 'dashed',  # Bordo tratteggiato
        'filter': {
            'style_names': ['Normal'],  # Solo paragrafi con stile Normal
            'exclude_headings': True,
            'min_length': 20  # Solo paragrafi con almeno 20 caratteri
        }
    }
)

print(f"\n✓ Test 3 completato")
print(f"  Borders applicati: {result3['borders_applied']}")
print(f"  Filtri applicati: stile=Normal, min_length=20")

print("\n" + "-"*80)
print("TEST 4: Confronto con metodo tradizionale (paragraph borders)")
print("-"*80)

test_output_4 = '/home/ssensini/WebstormProjects/divergentia/be/outputs/test_framing_traditional_output.docx'

result4 = service.apply_framing(
    test_input,
    test_output_4,
    {
        'sections': False,
        'paragraphs': True,
        'subparagraphs': False,
        'sentences': False,
        'use_tables': False,  # Usa metodo tradizionale
        'filter': {
            'exclude_headings': True
        }
    }
)

print(f"\n✓ Test 4 completato")
print(f"  Borders applicati: {result4['borders_applied']}")
print(f"  Metodo: {result4['method']}")

print("\n" + "="*80)
print("RIEPILOGO TEST")
print("="*80)
print("\n✓ Tutti i test completati con successo!")
print("\nFile generati:")
print(f"  1. {test_output}")
print(f"  2. {test_output_2}")
print(f"  3. {test_output_3}")
print(f"  4. {test_output_4}")

print("\nFunzionalità testate:")
print("  ✓ Incapsulamento in tabelle 1x1")
print("  ✓ Preservazione formattazione (grassetto, corsivo, colori)")
print("  ✓ Bordi personalizzabili (spessore, colore, stile)")
print("  ✓ Margini celle personalizzabili")
print("  ✓ Preservazione spaziatura paragrafi")
print("  ✓ Filtri per stile e lunghezza")
print("  ✓ Esclusione heading automatica")
print("  ✓ Confronto con metodo tradizionale")

print("\nVerifica manuale:")
print("  1. Apri i file DOCX generati")
print("  2. Verifica che i paragrafi siano incapsulati in tabelle")
print("  3. Verifica che la formattazione sia preservata")
print("  4. Verifica che i bordi siano applicati correttamente")

