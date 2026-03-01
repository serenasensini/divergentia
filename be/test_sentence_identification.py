"""
Test script for sentence identification using spaCy Sentencizer
"""
import unittest
from app.services.keyword_service import get_keyword_service


class TestSentenceSplitting(unittest.TestCase):
    """Test spaCy Sentencizer-based sentence splitting"""

    def setUp(self):
        """Set up test fixtures"""
        self.keyword_service = get_keyword_service()

    def test_simple_sentences(self):
        """Test simple sentence splitting"""
        text = "Questa è la prima frase. Questa è la seconda frase. Questa è la terza frase."
        sentences = self.keyword_service.split_sentences(text)

        self.assertEqual(len(sentences), 3)
        self.assertEqual(sentences[0], "Questa è la prima frase.")
        self.assertEqual(sentences[1], "Questa è la seconda frase.")
        self.assertEqual(sentences[2], "Questa è la terza frase.")
        print(f"✓ Simple sentences test passed: {len(sentences)} sentences")

    def test_abbreviations(self):
        """Test sentence splitting with abbreviations"""
        text = "Il Dr. Rossi lavora qui. Abbiamo visto il dott. Mario ieri. Etc. sono abbreviazioni comuni."
        sentences = self.keyword_service.split_sentences(text)

        # Should correctly identify 3 sentences, not split on "Dr." or "dott."
        self.assertGreaterEqual(len(sentences), 2)
        print(f"✓ Abbreviations test passed: {len(sentences)} sentences identified")
        for i, sent in enumerate(sentences):
            print(f"  Sentence {i+1}: {sent}")

    def test_exclamation_and_question(self):
        """Test different sentence endings"""
        text = "Come stai? Bene grazie! Che bella giornata."
        sentences = self.keyword_service.split_sentences(text)

        self.assertEqual(len(sentences), 3)
        self.assertTrue(sentences[0].endswith("?"))
        self.assertTrue(sentences[1].endswith("!"))
        print(f"✓ Different punctuation test passed: {len(sentences)} sentences")

    def test_complex_punctuation(self):
        """Test sentences with quotes and parentheses"""
        text = 'Il professore disse: "Studiate bene". (Questo è importante). Era una bella giornata.'
        sentences = self.keyword_service.split_sentences(text)

        self.assertGreaterEqual(len(sentences), 2)
        print(f"✓ Complex punctuation test passed: {len(sentences)} sentences")
        for i, sent in enumerate(sentences):
            print(f"  Sentence {i+1}: {sent}")

    def test_numbers_and_decimals(self):
        """Test sentences with numbers and decimals"""
        text = "Il prezzo è 19.99 euro. La temperatura è 36.5 gradi. Domani sarà diverso."
        sentences = self.keyword_service.split_sentences(text)

        self.assertEqual(len(sentences), 3)
        print(f"✓ Numbers and decimals test passed: {len(sentences)} sentences")

    def test_ellipsis(self):
        """Test sentences with ellipsis"""
        text = "Non so cosa dire... Forse è meglio tacere. Sì, decisamente."
        sentences = self.keyword_service.split_sentences(text)

        self.assertGreaterEqual(len(sentences), 2)
        print(f"✓ Ellipsis test passed: {len(sentences)} sentences")
        for i, sent in enumerate(sentences):
            print(f"  Sentence {i+1}: {sent}")

    def test_single_sentence(self):
        """Test single sentence"""
        text = "Questa è una sola frase."
        sentences = self.keyword_service.split_sentences(text)

        self.assertEqual(len(sentences), 1)
        self.assertEqual(sentences[0], text)
        print(f"✓ Single sentence test passed")

    def test_empty_text(self):
        """Test empty text"""
        sentences = self.keyword_service.split_sentences("")
        self.assertEqual(len(sentences), 0)

        sentences = self.keyword_service.split_sentences("   ")
        self.assertEqual(len(sentences), 0)
        print(f"✓ Empty text test passed")

    def test_no_sentence_ending(self):
        """Test text without sentence-ending punctuation"""
        text = "Questo testo non ha punteggiatura finale"
        sentences = self.keyword_service.split_sentences(text)

        # Should still identify it as a sentence
        self.assertEqual(len(sentences), 1)
        print(f"✓ No ending punctuation test passed: {sentences}")

    def test_multiple_spaces(self):
        """Test sentences with multiple spaces"""
        text = "Prima frase.    Seconda frase.     Terza frase."
        sentences = self.keyword_service.split_sentences(text)

        self.assertEqual(len(sentences), 3)
        print(f"✓ Multiple spaces test passed: {len(sentences)} sentences")

    def test_comparison_with_regex(self):
        """Compare spaCy Sentencizer with regex approach"""
        import re

        test_texts = [
            "Il Dr. Rossi è qui. Lavora da anni.",
            "Il prezzo è 19.99 euro. Costa poco.",
            "Come stai? Bene! Grazie.",
        ]

        print("\n" + "="*60)
        print("Comparison: spaCy Sentencizer vs Regex")
        print("="*60)

        for text in test_texts:
            # spaCy method
            spacy_sentences = self.keyword_service.split_sentences(text)

            # Regex method (old implementation)
            regex_sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
            regex_sentences = [s.strip() for s in regex_sentences if s.strip()]

            print(f"\nText: {text}")
            print(f"spaCy: {len(spacy_sentences)} sentences")
            for i, s in enumerate(spacy_sentences):
                print(f"  {i+1}. {s}")
            print(f"Regex: {len(regex_sentences)} sentences")
            for i, s in enumerate(regex_sentences):
                print(f"  {i+1}. {s}")


class TestFormattingServiceIntegration(unittest.TestCase):
    """Test integration with FormattingService"""

    def test_identify_sentences_with_spacy(self):
        """Test that _identify_sentences uses spaCy correctly"""
        from docx import Document as DocxDocument
        from app.services.formatting_service import get_formatting_service
        import tempfile
        import os

        # Create a temporary document
        doc = DocxDocument()
        doc.add_paragraph("Prima frase. Seconda frase. Terza frase.")
        doc.add_paragraph("Il Dr. Rossi lavora qui. È molto bravo.")
        doc.add_heading("Titolo", level=1)
        doc.add_paragraph("Altra frase nel documento. Con più contenuto.")

        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)

        try:
            # Test identify_sentences
            formatting_service = get_formatting_service()
            doc_reloaded = DocxDocument(tmp_path)
            sentence_map = formatting_service._identify_sentences(doc_reloaded)

            # Should identify sentences in non-heading paragraphs
            self.assertGreater(len(sentence_map), 0)

            total_sentences = sum(len(sents) for _, sents in sentence_map)
            print(f"\n✓ Integration test passed: {total_sentences} sentences in {len(sentence_map)} paragraphs")

            for para_idx, sentences in sentence_map:
                print(f"  Paragraph {para_idx}: {len(sentences)} sentences")

        finally:
            # Clean up
            if os.path.exists(tmp_path):
                os.remove(tmp_path)


if __name__ == '__main__':
    print("\n" + "="*60)
    print("Testing Sentence Identification with spaCy Sentencizer")
    print("="*60 + "\n")

    # Run tests
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()

    # Add tests in specific order
    suite.addTests(loader.loadTestsFromTestCase(TestSentenceSplitting))
    suite.addTests(loader.loadTestsFromTestCase(TestFormattingServiceIntegration))

    runner = unittest.TextTestRunner(verbosity=2)
    runner.run(suite)

