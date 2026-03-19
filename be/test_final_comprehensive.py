"""
Final comprehensive test for citations and headings preservation
"""
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.services.formatting_service import FormattingService

print("="*80)
print("FINAL TEST: Citations AND Headings Preservation")
print("="*80)

# Create test document
doc = Document()

# Add a "Normal" style paragraph that looks like a heading (no actual Heading style)
title_para = doc.add_paragraph("Le origini dell'interesse e della ricerca")
title_para.style = doc.styles['Normal']

# Add paragraph with text and a citation field
para = doc.add_paragraph()

# Add text before citation
run1 = para.add_run("A partire dalla fine del XIX secolo lo studio delle capacit intellettive")
run1.font.color.rgb = RGBColor(0, 0, 0)

# Add a run with fldChar to simulate a citation - THIS SHOULD BE PRESERVED
run_citation = para.add_run(" (Galton, 1869) ")
run_citation.font.color.rgb = RGBColor(128, 0, 128)  # Purple
fld_char = OxmlElement('w:fldChar')
fld_char.set(qn('w:fldCharType'), 'begin')
run_citation._element.append(fld_char)

# Add text after citation
run2 = para.add_run("diventato pi sistematico per merito di pioneiri.")
run2.font.color.rgb = RGBColor(0, 0, 0)

test_input = '/home/ssensini/WebstormProjects/divergentia/be/outputs/test_final_citation.docx'
test_output = '/home/ssensini/WebstormProjects/divergentia/be/outputs/test_final_citation_output.docx'
doc.save(test_input)

print("\nOriginal document:")
print(f"  Paragraph 0 (title): '{title_para.text}'")
print(f"    Style: {title_para.style.name}")
print(f"    Runs: {len(title_para.runs)}")

print(f"  Paragraph 1 (content): '{para.text[:80]}'")
print(f"    Runs: {len(para.runs)}")
for i, run in enumerate(para.runs):
    has_field = any('fldChar' in str(c.tag) for c in run._element if run._element)
    field_marker = " [HAS CITATION FIELD]" if has_field else ""
    print(f"      Run {i}: '{run.text}'{field_marker}")

print("\nApplying highlighting for NOUNS...")
print("-"*80)

formatting_service = FormattingService()

try:
    result = formatting_service.apply_highlighting(
        test_input,
        test_output,
        {
            'enabled': True,
            'color': '#FF0000',
            'style': 'bold',
            'font_size': None,
            'font_family': None,
            'nouns': True,
            'verbs': False,
            'adjectives': False,
            'adverbs': False
        }
    )

    print(f"✓ Highlighting applied")
    print(f"  Words formatted: {result['words_formatted']}")
    print(f"  Paragraphs processed: {result['paragraphs_processed']}")

    # Analyze output
    output_doc = Document(test_output)

    print("\nOutput document:")
    print(f"  Paragraph 0 (title): '{output_doc.paragraphs[0].text}'")
    print(f"    Runs: {len(output_doc.paragraphs[0].runs)}")

    if len(output_doc.paragraphs[0].runs) == 1:
        print("    ✓ Title was NOT processed (1 run - correct!)")
    else:
        print(f"    ✗ Title WAS processed ({len(output_doc.paragraphs[0].runs)} runs - WRONG!)")
        for i, run in enumerate(output_doc.paragraphs[0].runs[:5]):
            print(f"        Run {i}: '{run.text}'")

    print(f"  Paragraph 1 (content): '{output_doc.paragraphs[1].text[:80]}'")
    print(f"    Runs: {len(output_doc.paragraphs[1].runs)}")

    # Check citation position
    citation_found = False
    citation_position = -1

    for i, run in enumerate(output_doc.paragraphs[1].runs):
        has_field = any('fldChar' in str(c.tag) for c in run._element if run._element)
        if 'Galton' in run.text or has_field:
            citation_found = True
            citation_position = i
            color_str = "PURPLE" if run.font.color.rgb == RGBColor(128, 0, 128) else \
                       "RED" if run.font.color.rgb == RGBColor(255, 0, 0) else "OTHER"
            field_marker = " [HAS CITATION FIELD]" if has_field else ""
            print(f"      >>> Run {i}: '{run.text}' - {color_str}{field_marker}")

    if citation_found:
        if citation_position == 0:
            print(f"    ✗ Citation at BEGINNING (position {citation_position}) - WRONG!")
        elif citation_position == len(output_doc.paragraphs[1].runs) - 1:
            print(f"    ✗ Citation at END (position {citation_position}) - WRONG!")
        else:
            print(f"    ✓ Citation in MIDDLE (position {citation_position}/{len(output_doc.paragraphs[1].runs)}) - CORRECT!")
    else:
        print("    ✗ Citation NOT FOUND - WRONG!")

    print("\n" + "="*80)
    print("SUMMARY")
    print("="*80)
    print("Expected behavior:")
    print("  1. Title (paragraph 0) should have 1 run (not processed)")
    print("  2. Citation should be in the middle of paragraph 1")
    print("  3. Citation should keep purple color and field marker")
    print("  4. Only nouns (not the citation) should be red and bold")

except Exception as e:
    print(f"\n✗ Error: {str(e)}")
    import traceback
    traceback.print_exc()

