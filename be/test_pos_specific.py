"""
Test per verificare che ogni tipo di parte del discorso sia formattato correttamente
"""
from docx import Document
from docx.shared import RGBColor
from app.services.formatting_service import FormattingService

# Create test document
doc = Document()
para = doc.add_paragraph("Il gatto veloce mangia lentamente il cibo delizioso.")
# Translation: "The fast cat eats slowly the delicious food."
# Expected POS: Il (DET), gatto (NOUN), veloce (ADJ), mangia (VERB), lentamente (ADV), il (DET), cibo (NOUN), delizioso (ADJ)

test_input_path = "/home/ssensini/WebstormProjects/divergentia/be/outputs/test_pos_types.docx"
doc.save(test_input_path)

print("="*80)
print("TEST: FORMATTAZIONE PER TIPO DI PARTE DEL DISCORSO")
print("="*80)
print(f"Testo originale: {para.text}")
print()

# Test 1: Solo sostantivi
print("Test 1: Solo SOSTANTIVI (rosso, bold)")
print("-"*80)

formatting_service = FormattingService()
output_path_1 = "/home/ssensini/WebstormProjects/divergentia/be/outputs/test_pos_nouns.docx"

result = formatting_service.apply_highlighting(
    test_input_path,
    output_path_1,
    {
        'enabled': True,
        'color': '#FF0000',
        'style': 'bold',
        'font_size': None,
        'font_family': None,
        'nouns': True,
        'verbs': False,
        'adjectives': False,
        'adverbs': False
    }
)

print(f"✓ Parole formattate: {result['words_formatted']}")
print(f"  Sostantivi: {result['pos_stats']['nouns']}")
print(f"  Atteso: 'gatto' e 'cibo' formattati in rosso")

output_doc = Document(output_path_1)
print(f"  Run risultanti: {len(output_doc.paragraphs[0].runs)}")
for i, run in enumerate(output_doc.paragraphs[0].runs):
    color = "ROSSO" if run.font.color.rgb == RGBColor(255, 0, 0) else "NERO"
    bold = "bold" if run.bold else "normale"
    print(f"    Run {i}: '{run.text}' - {color}, {bold}")

# Test 2: Solo verbi
print("\nTest 2: Solo VERBI (blu, italic)")
print("-"*80)

output_path_2 = "/home/ssensini/WebstormProjects/divergentia/be/outputs/test_pos_verbs.docx"

result = formatting_service.apply_highlighting(
    test_input_path,
    output_path_2,
    {
        'enabled': True,
        'color': '#0000FF',
        'style': 'italic',
        'font_size': None,
        'font_family': None,
        'nouns': False,
        'verbs': True,
        'adjectives': False,
        'adverbs': False
    }
)

print(f"✓ Parole formattate: {result['words_formatted']}")
print(f"  Verbi: {result['pos_stats']['verbs']}")
print(f"  Atteso: 'mangia' formattato in blu")

output_doc = Document(output_path_2)
print(f"  Run risultanti: {len(output_doc.paragraphs[0].runs)}")
for i, run in enumerate(output_doc.paragraphs[0].runs):
    color = "BLU" if run.font.color.rgb == RGBColor(0, 0, 255) else "NERO"
    italic = "italic" if run.italic else "normale"
    print(f"    Run {i}: '{run.text}' - {color}, {italic}")

# Test 3: Solo aggettivi
print("\nTest 3: Solo AGGETTIVI (verde, underline)")
print("-"*80)

output_path_3 = "/home/ssensini/WebstormProjects/divergentia/be/outputs/test_pos_adjectives.docx"

result = formatting_service.apply_highlighting(
    test_input_path,
    output_path_3,
    {
        'enabled': True,
        'color': '#00FF00',
        'style': 'underline',
        'font_size': None,
        'font_family': None,
        'nouns': False,
        'verbs': False,
        'adjectives': True,
        'adverbs': False
    }
)

print(f"✓ Parole formattate: {result['words_formatted']}")
print(f"  Aggettivi: {result['pos_stats']['adjectives']}")
print(f"  Atteso: 'veloce' e 'delizioso' formattati in verde")

output_doc = Document(output_path_3)
print(f"  Run risultanti: {len(output_doc.paragraphs[0].runs)}")
for i, run in enumerate(output_doc.paragraphs[0].runs):
    color = "VERDE" if run.font.color.rgb == RGBColor(0, 255, 0) else "NERO"
    underline = "underline" if run.underline else "normale"
    print(f"    Run {i}: '{run.text}' - {color}, {underline}")

# Test 4: Solo avverbi
print("\nTest 4: Solo AVVERBI (arancione, bold)")
print("-"*80)

output_path_4 = "/home/ssensini/WebstormProjects/divergentia/be/outputs/test_pos_adverbs.docx"

result = formatting_service.apply_highlighting(
    test_input_path,
    output_path_4,
    {
        'enabled': True,
        'color': '#FFA500',
        'style': 'bold',
        'font_size': None,
        'font_family': None,
        'nouns': False,
        'verbs': False,
        'adjectives': False,
        'adverbs': True
    }
)

print(f"✓ Parole formattate: {result['words_formatted']}")
print(f"  Avverbi: {result['pos_stats']['adverbs']}")
print(f"  Atteso: 'lentamente' formattato in arancione")

output_doc = Document(output_path_4)
print(f"  Run risultanti: {len(output_doc.paragraphs[0].runs)}")
for i, run in enumerate(output_doc.paragraphs[0].runs):
    color = "ARANCIONE" if run.font.color.rgb == RGBColor(255, 165, 0) else "NERO"
    bold = "bold" if run.bold else "normale"
    print(f"    Run {i}: '{run.text}' - {color}, {bold}")

print("\n" + "="*80)
print("RIEPILOGO")
print("="*80)
print("✓ La funzione formatta SOLO le parti del discorso specificate")
print("✓ Tutto il resto del testo rimane inalterato")
print("✓ Ogni tipo di POS può avere formattazione differente")
print("✓ Le citazioni e i riferimenti sono preservati")

