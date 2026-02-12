"""
Text Extraction Utilities
"""
import logging
from pathlib import Path
from typing import Optional

from docx import Document
import PyPDF2
import pdfplumber

from app.exceptions.custom_exceptions import FileProcessingException

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text content from various file formats.

    Args:
        file_path: Path to file

    Returns:
        Extracted text content

    Raises:
        FileProcessingException: If extraction fails
    """
    file_extension = Path(file_path).suffix.lower().lstrip('.')

    logger.info(f"Extracting text from {file_extension} file: {file_path}")

    try:
        if file_extension == 'docx':
            return extract_text_from_docx(file_path)
        elif file_extension == 'pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension in ['txt', 'rtf']:
            return extract_text_from_txt(file_path)
        else:
            raise FileProcessingException(
                f"Unsupported file format for text extraction: {file_extension}"
            )
    except Exception as e:
        logger.error(f"Text extraction failed: {str(e)}")
        raise FileProcessingException(f"Failed to extract text: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted text
    """
    logger.debug(f"Extracting text from DOCX: {file_path}")

    doc = Document(file_path)
    text_parts = []

    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    full_text = '\n'.join(text_parts)
    logger.info(f"Extracted {len(full_text)} characters from DOCX")

    return full_text


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file.

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted text
    """
    logger.debug(f"Extracting text from PDF: {file_path}")

    text_parts = []

    try:
        # Try with pdfplumber first (better text extraction)
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
    except Exception as e:
        logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")

        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except Exception as e2:
            logger.error(f"PDF extraction failed: {str(e2)}")
            raise FileProcessingException(f"Failed to extract text from PDF: {str(e2)}")

    full_text = '\n'.join(text_parts)
    logger.info(f"Extracted {len(full_text)} characters from PDF")

    return full_text


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from TXT/RTF file.

    Args:
        file_path: Path to text file

    Returns:
        Extracted text
    """
    logger.debug(f"Reading text file: {file_path}")

    # Try different encodings
    encodings = ['utf-8', 'latin-1', 'cp1252']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()

            logger.info(f"Read {len(text)} characters from text file (encoding: {encoding})")
            return text
        except UnicodeDecodeError:
            continue

    raise FileProcessingException(f"Failed to read text file with supported encodings")


def count_words(text: str) -> int:
    """
    Count words in text.

    Args:
        text: Text to count

    Returns:
        Word count
    """
    return len(text.split())


def count_sentences(text: str) -> int:
    """
    Count sentences in text (approximate).

    Args:
        text: Text to analyze

    Returns:
        Sentence count (approximate)
    """
    # Simple sentence counting
    sentence_endings = ['. ', '! ', '? ']
    count = 0
    for ending in sentence_endings:
        count += text.count(ending)
    return max(1, count)


def get_text_preview(text: str, length: int = 200) -> str:
    """
    Get preview of text.

    Args:
        text: Full text
        length: Preview length in characters

    Returns:
        Text preview
    """
    if len(text) <= length:
        return text

    return text[:length].strip() + '...'
