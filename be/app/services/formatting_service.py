"""
Document Formatting Service - Handle document style modifications
"""
import logging
import re
from typing import Dict, Any, Optional, List, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
# import pymupdf  # PyMuPDF

from app.exceptions.custom_exceptions import (
    FormattingException,
)
from app.services.document_structure_service import (
    BlockType,
    DocumentStructure,
    get_document_structure_service,
)
from app.utils.color_utils import DEFAULT_SCHEME, generate_palette

logger = logging.getLogger(__name__)


class FormattingService:
    """Service for applying formatting changes to documents"""

    SUPPORTED_FORMATS = {
        'docx': ['font_name', 'font_size', 'font_color', 'bold', 'italic', 'alignment', 'framing'],
        'pdf': [],  # Limited PDF support
        'txt': []  # Plain text has no formatting
    }

    # Heading identification thresholds
    HEADING_FONT_SIZE_THRESHOLD = 14.0  # Font size in points to consider as heading
    MIN_WORD_LENGTH = 2  # Minimum word length for keyword extraction

    # Default color specifications
    DEFAULT_HIGHLIGHT_COLOR = "#cc5500"  # Default highlight color (orange)
    DEFAULT_TEXT_COLOR_RGB = (0, 0, 0)  # Black RGB color tuple

    # Default border specifications
    DEFAULT_BORDER_WIDTH = 4  # 1/2 pt = 4 eighths of a point
    DEFAULT_BORDER_COLOR = "000000"  # Black (hex without #)
    DEFAULT_BORDER_STYLE = "single"  # Solid line

    # Table-based framing specifications
    DEFAULT_TABLE_BORDER_WIDTH = 8  # 1 pt = 8 eighths of a point
    DEFAULT_TABLE_CELL_MARGIN = 100  # Twips (1/1440 inch) - approx 1.76mm
    DEFAULT_TABLE_SPACING = 120  # Twips - spacing between table boxes

    # Per-part default border styles/widths, so sections, paragraphs and
    # sentences are visually distinguishable when framed with tables. Widths are
    # in eighths of a point. These apply only when the request does not specify
    # an explicit border_style / border_width.
    SECTION_BORDER_STYLE = "double"
    SECTION_BORDER_WIDTH = 16   # 2 pt
    PARAGRAPH_BORDER_STYLE = "single"
    PARAGRAPH_BORDER_WIDTH = 8  # 1 pt
    SENTENCE_BORDER_STYLE = "dashed"
    SENTENCE_BORDER_WIDTH = 4   # 1/2 pt
    
    # Font size validation
    MIN_FONT_SIZE = 6  # Minimum font size in points
    MAX_FONT_SIZE = 72  # Maximum font size in points

    # Spacing for list items (ordered/unordered). Instead of inserting blank
    # lines, list entries get a padding above and below each element. ~10px at
    # 96 dpi (1px = 0.75pt) => 7.5pt of space before/after.
    LIST_ITEM_SPACING_PT = 7.5

    # Space added after an inserted keywords paragraph so the keywords are
    # separated from the following section content. Twips (1pt = 20 twips) =>
    # 200 twips = 10pt.
    KEYWORD_SPACING_AFTER_TWIPS = 200

    def __init__(self) -> None:
        """Initialize formatting service"""
        self._structure_service = get_document_structure_service()
        logger.info("Formatting service initialized")

    def _build_structure(self, doc: Document, source_path: Optional[str] = None) -> DocumentStructure:
        """
        Build the canonical structural model of a document.

        The structure is computed once per operation and passed to the various
        ``_identify_*`` helpers, so heading/section/paragraph detection is
        consistent and based on style outline levels (with a Markdown fallback).

        Args:
            doc: python-docx Document object
            source_path: Original .docx path (enables the Markdown fallback)

        Returns:
            DocumentStructure describing every paragraph
        """
        return self._structure_service.build(doc, source_path=source_path)

    def _is_heading(self, para, check_font_size: bool = True, font_size_threshold: float = None) -> bool:
        """
        Check if a paragraph is a heading (section delimiter).

        Classification is delegated to the DocumentStructureService: the
        paragraph style outline level is used first, and only when the style
        carries no heading information (and ``check_font_size`` is True) a
        font-size/bold heuristic is applied.

        Args:
            para: python-docx paragraph object
            check_font_size: Whether to allow the heuristic fallback
            font_size_threshold: Unused, kept for backward compatibility

        Returns:
            bool: True if paragraph is a heading, False otherwise
        """
        block = self._structure_service.classify_paragraph(para, use_heuristic=check_font_size)
        return block.is_heading

    def _is_main_title(self, para) -> bool:
        """
        Check if a paragraph is a main/document title (Title style or Heading 1).

        Args:
            para: python-docx paragraph object

        Returns:
            bool: True if paragraph is a main title, False otherwise
        """
        return self._structure_service.classify_paragraph(para).type == BlockType.DOCUMENT_TITLE

    def _is_section_heading(self, para) -> bool:
        """
        Check if a paragraph is a section heading (Heading 2, 3, 4, ...).

        Args:
            para: python-docx paragraph object

        Returns:
            bool: True if paragraph is a section heading, False otherwise
        """
        return self._structure_service.classify_paragraph(para).type == BlockType.SECTION_HEADING

    def _get_paragraph_font_size(self, para) -> Optional[float]:
        """
        Get the font size of a paragraph in points.

        Checks both the paragraph style and individual runs to find the font size.

        Args:
            para: python-docx paragraph object

        Returns:
            Optional[float]: Font size in points, or None if not found
        """
        try:
            # First try to get font size from style
            if para.style.font.size:
                logger.debug(f"Paragraph style font size found: {para.style.font.size.pt} pt")
                return para.style.font.size.pt
        except (AttributeError, TypeError):
            pass

        # Fall back to checking runs
        for run in para.runs:
            try:
                if run.font.size:
                    logger.debug(f"Run font size found: {run.font.size.pt} pt")
                    return run.font.size.pt
            except (AttributeError, TypeError):
                continue

        return None

    def apply_formatting(
        self,
        input_path: str,
        output_path: str,
        formatting_options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Apply formatting changes to a document.

        Args:
            input_path: Path to input document
            output_path: Path to save formatted document
            formatting_options: Dictionary with formatting options

        Returns:
            Dictionary with result information

        Raises:
            FormattingException: If formatting fails
        """
        logger.info(f"Applying formatting to {input_path}")

        # Determine file type
        file_extension = Path(input_path).suffix.lower().lstrip('.')

        if file_extension not in self.SUPPORTED_FORMATS:
            raise FormattingException(
                f"Unsupported file format: {file_extension}",
                payload={'supported_formats': list(self.SUPPORTED_FORMATS.keys())}
            )

        # Apply formatting based on file type
        try:
            if file_extension == 'docx':
                return self._format_docx(input_path, output_path, formatting_options)
            elif file_extension == 'pdf':
                return self._format_pdf(input_path, output_path, formatting_options)
            else:
                raise FormattingException(f"Handler not implemented for {file_extension}")

        except Exception as e:
            logger.error(f"Formatting error: {str(e)}")
            raise FormattingException(f"Failed to apply formatting: {str(e)}")

    def apply_framing(
        self,
        input_path: str,
        output_path: str,
        framing_options: Dict[str, bool]
    ) -> dict[str, Any] | str:
        """
        Apply framing (borders) to document parts.

        Args:
            input_path: Path to input document
            output_path: Path to save framed document
            framing_options: Dictionary with boolean flags for each document part

        Returns:
            Dictionary with result information

        Raises:
            FormattingException: If framing fails
        """
        logger.info(f"Applying framing to {input_path}")

        # Determine file type
        file_extension = Path(input_path).suffix.lower().lstrip('.')

        if file_extension not in self.SUPPORTED_FORMATS:
            raise FormattingException(
                f"Unsupported file format: {file_extension}",
                payload={'supported_formats': list(self.SUPPORTED_FORMATS.keys())}
            )

        # Apply framing based on file type
        try:
            if file_extension == 'docx':
                logger.debug(f"Framing options received: {framing_options}")
                return self._apply_framing_docx(input_path, output_path, framing_options)
            elif file_extension == 'pdf':
                return "PDF framing is currently under development and has limited support"
            else:
                raise FormattingException(f"Framing not supported for {file_extension}")

        except Exception as e:
            logger.error(f"Framing error: {str(e)}")
            raise FormattingException(f"Failed to apply framing: {str(e)}")

    def apply_spacing(
        self,
        input_path: str,
        output_path: str,
        spacing_options: Dict[str, bool]
    ) -> Dict[str, Any]:
        """
        Apply spacing to document parts.

        Args:
            input_path: Path to input document
            output_path: Path to save framed document
            spacing_options: Dictionary with boolean flags for each document part

        Returns:
            Dictionary with result information

        Raises:
            FormattingException: If framing fails
        """
        logger.info(f"Applying spacing to {input_path}")

        # Determine file type
        file_extension = Path(input_path).suffix.lower().lstrip('.')

        if file_extension not in self.SUPPORTED_FORMATS:
            raise FormattingException(
                f"Unsupported file format: {file_extension}",
                payload={'supported_formats': list(self.SUPPORTED_FORMATS.keys())}
            )

        # Apply framing based on file type
        try:
            if file_extension == 'docx':
                logger.debug(f"Spacing options received: {spacing_options}")
                return self._apply_spacing_docx(input_path, output_path, spacing_options)
            elif file_extension == 'pdf':
                return "PDF spacing is currently under development and has limited support"
                # return self._apply_framing_pdf(input_path, output_path, framing_options)
            else:
                raise FormattingException(f"Framing not supported for {file_extension}")

        except Exception as e:
            logger.error(f"Framing error: {str(e)}")
            raise FormattingException(f"Failed to apply framing: {str(e)}")

    def apply_keywords(
        self,
        input_path: str,
        output_path: str,
        keyword_options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Extract keywords from document sections and add them as initial paragraphs.

        Args:
            input_path: Input file path
            output_path: Output file path
            keyword_options: Dictionary with keyword extraction options

        Returns:
            Result information including number of sections processed
        """
        logger.info(f"Applying keyword extraction to {input_path}")

        # Determine file type
        file_extension = Path(input_path).suffix.lower().lstrip('.')

        if file_extension not in self.SUPPORTED_FORMATS:
            raise FormattingException(
                f"Unsupported file format: {file_extension}",
                payload={'supported_formats': list(self.SUPPORTED_FORMATS.keys())}
            )

        # Apply keywords based on file type
        try:
            if file_extension == 'docx':
                logger.debug(f"Keyword options received: {keyword_options}")
                return self._apply_keywords_docx(input_path, output_path, keyword_options)
            elif file_extension == 'pdf':
                raise FormattingException("PDF keyword extraction is not currently supported")
            else:
                raise FormattingException(f"Keyword extraction not supported for {file_extension}")

        except Exception as e:
            logger.error(f"Keyword extraction error: {str(e)}")
            raise FormattingException(f"Failed to apply keyword extraction: {str(e)}")

    def apply_highlighting(
        self,
        input_path: str,
        output_path: str,
        highlighting_options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Apply part-of-speech text formatting to document text.

        Args:
            input_path: Input file path
            output_path: Output file path
            highlighting_options: Dictionary with text formatting options
                - enabled: bool - Enable text formatting
                - color: str - Text color in hex format (e.g., "#FF0000")
                - style: str - Text styles: 'bold', 'italic', 'underline', or combinations like 'bold,italic'
                - font_size: int - Font size in points (6-72)
                - font_family: str - Font family name (e.g., 'Times New Roman', 'Arial', 'Courier New')
                - nouns: bool - Format nouns
                - verbs: bool - Format verbs
                - adjectives: bool - Format adjectives
                - adverbs: bool - Format adverbs

        Returns:
            Result information including number of words formatted
        """
        logger.info(f"Applying part-of-speech text formatting to {input_path}")

        # Determine file type
        file_extension = Path(input_path).suffix.lower().lstrip('.')

        if file_extension not in self.SUPPORTED_FORMATS:
            raise FormattingException(
                f"Unsupported file format: {file_extension}",
                payload={'supported_formats': list(self.SUPPORTED_FORMATS.keys())}
            )

        # Check if highlighting is enabled
        if not highlighting_options.get('enabled', False):
            raise FormattingException("Text formatting is not enabled in the provided options")

        # Validate at least one POS is selected
        pos_selected = any([
            highlighting_options.get('nouns', False),
            highlighting_options.get('verbs', False),
            highlighting_options.get('adjectives', False),
            highlighting_options.get('adverbs', False)
        ])

        if not pos_selected:
            raise FormattingException("At least one part of speech must be selected for text formatting")

        # Apply highlighting based on file type
        try:
            if file_extension == 'docx':
                logger.debug(f"Text formatting options received: {highlighting_options}")
                return self._apply_highlighting_docx(input_path, output_path, highlighting_options)
            elif file_extension == 'pdf':
                raise FormattingException("PDF text formatting is not currently supported")
            else:
                raise FormattingException(f"Text formatting not supported for {file_extension}")

        except Exception as e:
            logger.error(f"Text formatting error: {str(e)}")
            raise FormattingException(f"Failed to apply text formatting: {str(e)}")

    def _format_docx(
        self,
        input_path: str,
        output_path: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Format DOCX document.

        Args:
            input_path: Input file path
            output_path: Output file path
            options: Formatting options

            Example options:
            {
            "formatting": {
                "titles": false,
                "paragraphs": false,
                "section_titles": false,
                "paragraphs_titles": false,
                "captions": false,
                "bibliography": false,
                "theme": {
                  "positive": "#FF0000",
                  "negative": "#0000FF",
                  "scheme": "even"
                }
            }
        }

        The ``theme`` provides up to two seed colors (``positive`` /
        ``negative``). When more than two roles are colored at once, extra
        colors are derived from the seeds using the ``scheme`` harmony
        (complementary, triadic, tetradic, even, analogous) so every enabled
        role receives a distinct color.

        Returns:
            Result information
        """
        logger.info("Formatting DOCX document")

        try:
            doc = Document(input_path)
            paragraphs_modified = 0

            # Build the structural model once for this operation
            structure = self._build_structure(doc, source_path=input_path)
            logger.debug(f"Options received for formatting: {options}")

            theme = options.get('theme')
            logger.debug(f"Received theme option: {theme}")
            primary_col = theme.get('positive') if theme else None
            secondary_col = theme.get('negative') if theme else None
            scheme = (theme.get('scheme') if theme else None) or DEFAULT_SCHEME
            logger.debug(f"Primary color: {primary_col}, Secondary color: {secondary_col}, scheme: {scheme}")

            # Map options to their corresponding identification methods.
            # The order is canonical so palette assignment is deterministic.
            formatting_tasks = [
                ('titles', self._identify_main_title, 'main titles'),
                ('paragraphs', self._identify_paragraphs, 'paragraphs'),
                ('section_titles', self._identify_section_titles, 'section titles'),
                ('paragraphs_titles', self._identify_paragraph_titles, 'paragraph titles'),
                ('captions', self._identify_image_captions, 'image captions')
            ]

            # Build a harmonized palette sized to the number of enabled roles so
            # each colored role gets a distinct, related color (derived from the
            # user's positive/negative seeds via the chosen color-wheel scheme).
            enabled_tasks = [t for t in formatting_tasks if options.get(t[0])]
            palette = generate_palette([primary_col, secondary_col], len(enabled_tasks), scheme)
            logger.debug(f"Generated palette for {len(enabled_tasks)} roles: {palette}")

            for color, (option_key, identify_method, label) in zip(palette, enabled_tasks):
                indices = identify_method(doc, structure)
                logger.debug(f"Found {label}: {indices}")
                # Apply the palette color assigned to this role
                color_rgb = self._apply_color_to_indices(doc, indices, color, label)
                # Captions may also live inside text boxes anchored to images,
                # which are not part of doc.paragraphs: color those too.
                if option_key == 'captions':
                    self._color_textbox_captions(doc, color_rgb)

            if options.get('bibliography'):
                pass

            # Save document
            doc.save(output_path)

            logger.info(f"DOCX formatting completed: {paragraphs_modified} paragraphs modified")

            return {
                'success': True,
                'output_path': output_path,
                'format': 'docx',
                'paragraphs_modified': paragraphs_modified,
                'applied_options': options
            }

        except Exception as e:
            logger.error(f"DOCX formatting error: {str(e)}")
            raise FormattingException(f"DOCX formatting failed: {str(e)}")

    def _format_pdf(
        self,
        input_path: str,
        output_path: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Format PDF document (limited support - creates new PDF with styling).

        Note: PDF formatting is limited. This creates a new styled version.

        Args:
            input_path: Input file path
            output_path: Output file path
            options: Formatting options

        Returns:
            Result information
        """
        # logger.info("Formatting PDF document")
        # logger.warning("PDF formatting has limited support")
        #
        # try:
        #     # Open PDF
        #     doc = pymupdf.open(input_path)
        #     pages_processed = 0
        #
        #     # Extract formatting options
        #     font_name = options.get('font_name', 'helv')
        #     font_size = options.get('font_size', 11)
        #     font_color = options.get('font_color', '#000000')
        #
        #     # Parse color
        #     color_rgb = self._parse_color(font_color)
        #     color_tuple = tuple(c / 255.0 for c in color_rgb) if color_rgb else (0, 0, 0)
        #
        #     # Create new PDF with styling
        #     output_doc = pymupdf.open()
        #
        #     for page in doc:
        #         # Extract text
        #         text = page.get_text()
        #
        #         # Create new page
        #         new_page = output_doc.new_page(width=page.rect.width, height=page.rect.height)
        #
        #         # Insert text with new formatting
        #         rect = pymupdf.Rect(50, 50, page.rect.width - 50, page.rect.height - 50)
        #         new_page.insert_textbox(
        #             rect,
        #             text,
        #             fontsize=float(font_size),
        #             fontname=font_name,
        #             color=color_tuple
        #         )
        #
        #         pages_processed += 1
        #
        #     # Save output
        #     output_doc.save(output_path)
        #     output_doc.close()
        #     doc.close()
        #
        #     logger.info(f"PDF formatting completed: {pages_processed} pages processed")
        #
        #     return {
        #         'success': True,
        #         'output_path': output_path,
        #         'format': 'pdf',
        #         'pages_processed': pages_processed,
        #         'applied_options': options,
        #         'note': 'PDF formatting support is limited'
        #     }
        #
        # except Exception as e:
        #     # logger.error(f"PDF formatting error: {str(e)}")
        #     raise FormattingException(f"PDF formatting failed: {str(e)}")


    def _parse_color(self, color_str: str) -> Optional[tuple]:
        """
        Parse color string to RGB tuple.

        Args:
            color_str: Color string (hex format: #RRGGBB or RRGGBB)

        Returns:
            RGB tuple or None
        """
        try:
            # Remove # if present
            color_str = color_str.lstrip('#')

            # Parse hex to RGB
            r = int(color_str[0:2], 16)
            g = int(color_str[2:4], 16)
            b = int(color_str[4:6], 16)
            logger.debug(f"Parsed color {color_str} to RGB: ({r}, {g}, {b})")
            return (r, g, b)

        except (ValueError, IndexError) as e:
            logger.warning(f"Invalid color format: {color_str}")
            return None

    def _apply_color_to_indices(
        self,
        doc: Document,
        indices: List[int],
        color: Optional[str],
        label: str
    ) -> Optional[tuple]:
        """
        Apply a single color to paragraph runs at the specified indices.

        Args:
            doc: python-docx Document object
            indices: List of paragraph indices to color
            color: Hex color to apply (``#RRGGBB``), or None to skip
            label: Label for logging purposes

        Returns:
            The RGB tuple that was applied, or None if no color was available.
        """
        if not color:
            logger.warning(f"No color available for {label}")
            return None

        color_rgb = self._parse_color(color)
        logger.debug(f"Parsed color RGB for {label}: {color_rgb}")
        if color_rgb:
            for idx in indices:
                for run in doc.paragraphs[idx].runs:
                    run.font.color.rgb = RGBColor(*color_rgb)
        return color_rgb

    def _iter_textbox_paragraphs(self, doc: Document) -> List[Paragraph]:
        """
        Return the paragraphs contained in text boxes (``w:txbxContent``).

        Captions produced by Word's "Insert Caption" on a floating image are
        stored inside a text box anchored to the drawing, so they never appear
        in ``doc.paragraphs``. This helper exposes them as python-docx
        Paragraph objects so they can be inspected and formatted.
        """
        paragraphs: List[Paragraph] = []
        for txbx in doc.element.body.iter(qn('w:txbxContent')):
            for p_element in txbx.iter(qn('w:p')):
                paragraphs.append(Paragraph(p_element, doc))
        return paragraphs

    def _color_textbox_captions(self, doc: Document, color_rgb: Optional[tuple]) -> int:
        """
        Color caption paragraphs that live inside text boxes.

        Args:
            doc: python-docx Document object
            color_rgb: RGB tuple to apply (no-op if None)

        Returns:
            Number of textbox caption paragraphs colored.
        """
        if not color_rgb:
            return 0

        colored = 0
        for paragraph in self._iter_textbox_paragraphs(doc):
            block = self._structure_service.classify_paragraph(paragraph)
            if block.type == BlockType.CAPTION:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(*color_rgb)
                colored += 1
                logger.debug(f"Colored textbox caption: '{paragraph.text[:50]}...'")
        logger.info(f"Colored {colored} textbox caption paragraph(s)")
        return colored

    def _add_paragraph_border(self, paragraph) -> None:
        """
        Add border to a paragraph.

        Args:
            paragraph: python-docx paragraph object
        """
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), self.DEFAULT_BORDER_STYLE)
            border.set(qn('w:sz'), str(self.DEFAULT_BORDER_WIDTH))
            border.set(qn('w:space'), '1')
            border.set(qn('w:color'), self.DEFAULT_BORDER_COLOR)
            pBdr.append(border)

        pPr.append(pBdr)

    # FIXME: funziona con i paragrafi ma non con le frasi, da capire se è un problema di identificazione o di applicazione del bordo
    def _add_sentence_border(self, sentence) -> None:
        """
        Add border to a sentence.

        Args:
            paragraph: python-docx paragraph object
        """
        pPr = sentence._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), self.DEFAULT_BORDER_STYLE)
            border.set(qn('w:sz'), str(self.DEFAULT_BORDER_WIDTH))
            border.set(qn('w:space'), '1')
            border.set(qn('w:color'), self.DEFAULT_BORDER_COLOR)
            pBdr.append(border)

        pPr.append(pBdr)


    def _is_list_paragraph(self, paragraph) -> bool:
        """
        Detect whether a paragraph is a list item (ordered or unordered).

        Two signals are used:
        - Numbering properties (``w:numPr`` in the paragraph properties), which
          Word applies to both bulleted and numbered lists.
        - The paragraph style name (e.g. "List Paragraph", "List Bullet",
          "List Number" and localised variants such as "Paragrafo elenco").

        Args:
            paragraph: python-docx paragraph object

        Returns:
            True if the paragraph is a list element, False otherwise
        """
        pPr = paragraph._p.pPr
        if pPr is not None and pPr.find(qn('w:numPr')) is not None:
            return True

        style_name = (getattr(paragraph.style, 'name', '') or '').lower()
        list_markers = (
            'list paragraph', 'list bullet', 'list number', 'list continue',
            'paragrafo elenco', 'elenco puntato', 'elenco numerato',
        )
        return any(marker in style_name for marker in list_markers)

    def _apply_list_item_spacing(self, paragraph) -> None:
        """
        Add padding above and below a list item instead of blank lines.

        Keeps list items visually spaced without breaking the list structure or
        inserting stray line breaks between the bullet/number and its text.

        Word's "contextual spacing" (``w:contextualSpacing``) is disabled on the
        paragraph as well: list styles commonly enable it, which suppresses the
        space before/after *between* consecutive items of the same style and
        would otherwise make only the first element look spaced.

        Args:
            paragraph: python-docx paragraph object
        """
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(self.LIST_ITEM_SPACING_PT)
        paragraph_format.space_after = Pt(self.LIST_ITEM_SPACING_PT)

        pPr = paragraph._p.get_or_add_pPr()
        contextual_spacing = pPr.find(qn('w:contextualSpacing'))
        if contextual_spacing is None:
            contextual_spacing = OxmlElement('w:contextualSpacing')
            pPr.append(contextual_spacing)
        contextual_spacing.set(qn('w:val'), '0')

    def _add_paragraph_spacing(self, paragraph) -> None:
        """
        Add a new line before and after a paragraph.

        List items (ordered/unordered) are handled differently: instead of
        blank lines they receive padding above and below each element, so the
        list structure stays intact.

        Args:
            paragraph: python-docx paragraph object
        """
        if self._is_list_paragraph(paragraph):
            self._apply_list_item_spacing(paragraph)
            return

        # Add a new line before the input paragraph
        p = paragraph._element
        # Add <w:br/> before the paragraph
        p.addprevious(OxmlElement('w:br'))
        # Add <w:br/> after the paragraph
        p.addnext(OxmlElement('w:br'))

    def _add_sentence_spacing(self, paragraph) -> None:
        """
        Add spacing between sentences in a paragraph using spaCy Sentencizer.

        This method uses spaCy's Sentencizer for accurate sentence boundary detection,
        which properly handles abbreviations, complex punctuation, and language-specific rules.

        Args:
            paragraph: python-docx paragraph object
        """
        from app.services.keyword_service import get_keyword_service

        text = paragraph.text
        logger.debug(f"Adding sentence spacing to paragraph: '{text[:50]}...'")

        if not text.strip():
            return

        # List items keep their structure: apply padding around the element
        # instead of splitting the entry across lines.
        if self._is_list_paragraph(paragraph):
            logger.debug("Paragraph is a list item; applying padding instead of line breaks")
            self._apply_list_item_spacing(paragraph)
            return

        # Use spaCy Sentencizer for accurate sentence splitting
        keyword_service = get_keyword_service()
        sentences = keyword_service.split_sentences(text)

        logger.debug(f"Identified {len(sentences)} sentences for spacing")

        if len(sentences) <= 1:
            logger.debug("Only one sentence found, no spacing needed")
            return  # No need to add breaks if there's only one sentence

        # Clear existing runs while preserving paragraph properties
        for run in paragraph.runs:
            r = run._element
            r.getparent().remove(r)

        # Re-add sentences separated by a single line break. A single <w:br/>
        # puts each sentence on its own line without the doubled blank-line gap
        # that previously appeared between words/sentences.
        for i, sentence in enumerate(sentences):
            logger.debug(f"Adding sentence {i+1}/{len(sentences)}: '{sentence[:30]}...'")
            run = paragraph.add_run(sentence)

            # Add a single line break after each sentence except the last one
            if i < len(sentences) - 1:
                run.add_break()

        logger.debug(f"Finished adding sentence spacing to paragraph")


    def _identify_sections(self, doc: Document, structure: Optional[DocumentStructure] = None) -> List[Tuple[int, int, str]]:
        """
        Identifica le sezioni del documento. Una sezione è il testo compreso tra
        due heading (o da un heading fino alla fine del documento), heading esclusi.

        La classificazione degli heading è delegata al DocumentStructureService,
        che usa il livello di outline dello stile (Titolo / Heading 1 / Heading 2
        ...) con fallback a Markdown / euristica.

        Args:
            doc: python-docx Document object
            structure: struttura pre-calcolata (opzionale)

        Returns:
            Lista di tuple (start_index, end_index, section_text)
        """
        structure = structure or self._build_structure(doc)
        sections = [section.as_tuple() for section in structure.sections()]

        logger.info(f"Identified {len(sections)} sections")
        for idx, (start, end, text) in enumerate(sections):
            logger.debug(f"Section {idx}: paragraphs {start}-{end}, text preview: '{text[:100]}...'")

        return sections

    def _identify_paragraphs(self, doc: Document, structure: Optional[DocumentStructure] = None) -> List[int]:
        """
        Identify content paragraphs in the document.

        A paragraph is any non-empty paragraph that is not a heading, a document
        title or a caption. Content is captured regardless of the heading level
        it sits under (including the introductory part before/under top-level
        headings), not only under Heading 2+ sections.

        Args:
            doc: python-docx Document object
            structure: pre-computed structure (optional)

        Returns:
            List of paragraph indices that are content paragraphs
        """
        structure = structure or self._build_structure(doc)
        paragraphs = structure.content_paragraph_indices()

        logger.info(f"Identified {len(paragraphs)} paragraphs (text following section headings)")
        if paragraphs:
            logger.debug(f"Example of identified paragraphs: {[doc.paragraphs[i].text[:30] for i in paragraphs[:3]]}")
        return paragraphs

    def _identify_subparagraphs(self, doc: Document, structure: Optional[DocumentStructure] = None) -> List[Tuple[int, List[int]]]:
        """
        Identify subparagraphs within paragraphs.
        Subparagraphs are complex periods dependent on each other.

        Args:
            doc: python-docx Document object
            structure: pre-computed structure (optional, unused here)

        Returns:
            List of tuples (paragraph_index, [run_indices])
        """
        subparagraphs = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text
            if not text.strip():
                continue

            # Split by complex sentence markers (;, :, etc.)
            parts = re.split(r'[;:]', text)
            if len(parts) > 1:
                run_groups = list(range(len(parts)))
                subparagraphs.append((i, run_groups))

        logger.info(f"Identified {len(subparagraphs)} paragraphs with subparagraphs")
        return subparagraphs

    def _identify_sentences(self, doc: Document, structure: Optional[DocumentStructure] = None) -> List[Tuple[int, List[str]]]:
        """
        Identify sentences in document using spaCy's Sentencizer.

        Headings (identified via the structural model) are skipped, as they are
        not sentence content.

        Args:
            doc: python-docx Document object
            structure: pre-computed structure (optional)

        Returns:
            List of tuples (paragraph_index, [sentences])
        """
        from app.services.keyword_service import get_keyword_service

        structure = structure or self._build_structure(doc)
        sentence_map = []
        keyword_service = get_keyword_service()

        for i, para in enumerate(doc.paragraphs):
            # Skip headings as they're typically not sentence content
            if structure.is_heading(i):
                logger.debug(f"Skipping sentence identification for paragraph {i} (heading)")
                continue

            text = para.text
            if not text.strip():
                continue

            sentences = keyword_service.split_sentences(text)
            if sentences:
                sentence_map.append((i, sentences))

        total_sentences = sum(len(sents) for _, sents in sentence_map)
        logger.info(f"Identified {total_sentences} sentences across {len(sentence_map)} paragraphs using spaCy Sentencizer")

        return sentence_map

    # Identify main title (Title style / Heading 1) - special case for framing main title only
    def _identify_main_title(self, doc: Document, structure: Optional[DocumentStructure] = None) -> List[int]:
        """
        Identify main/document titles (Title style or Heading 1).

        Args:
            doc: python-docx Document object
            structure: pre-computed structure (optional)

        Returns:
            List of paragraph indices identified as main titles
        """
        structure = structure or self._build_structure(doc)
        title_indices = structure.main_title_indices()

        logger.info(f"Identified {len(title_indices)} titles")
        logger.debug(f"Identified title indices: {title_indices}")
        return title_indices

    # Identify section titles (Heading 1) - special case for framing section titles only
    def _identify_section_titles(self, doc: Document, structure: Optional[DocumentStructure] = None) -> List[int]:
        """
        Identify section headings (Heading 1 / outline level 0 / Markdown '#').

        Args:
            doc: python-docx Document object
            structure: pre-computed structure (optional)

        Returns:
            List of paragraph indices identified as section titles
        """
        structure = structure or self._build_structure(doc)
        title_indices = structure.section_heading_indices()

        logger.info(f"Identified {len(title_indices)} section titles")
        return title_indices

    # Identify paragraph titles (Heading 2 and deeper)
    def _identify_paragraph_titles(self, doc: Document, structure: Optional[DocumentStructure] = None) -> List[int]:
        """
        Identify paragraph titles (Heading 2, 3, ... / outline level >= 1 /
        Markdown '##'+).

        These are the sub-section labels sitting below section titles, kept
        distinct from both the document title and the section titles so they
        can be formatted independently.

        Args:
            doc: python-docx Document object
            structure: pre-computed structure (optional)

        Returns:
            List of paragraph indices identified as paragraph titles
        """
        structure = structure or self._build_structure(doc)
        title_indices = structure.paragraph_title_indices()

        logger.info(f"Identified {len(title_indices)} paragraph titles")
        return title_indices

    # Identify captions (figure/table captions)
    def _identify_image_captions(self, doc: Document, structure: Optional[DocumentStructure] = None) -> List[int]:
        """
        Identify captions in the document.

        A caption is detected from its paragraph style (localised names such as
        "Caption" / "Didascalia" / "Légende" ...) or from its textual pattern
        (e.g. "Figura 2: ...", "Figure 2. ...", "Tabella 1 - ...").

        Detection is intentionally independent of the presence of an adjacent
        image: captions styled or worded as such are recognised even when the
        related image is anchored/floating or separated by empty paragraphs.

        Args:
            doc: python-docx Document object
            structure: pre-computed structure (optional)
        Returns:
            List of paragraph indices that are identified as captions
        """
        structure = structure or self._build_structure(doc)
        caption_indices = structure.caption_indices()
        logger.info(f"Identified {len(caption_indices)} captions")
        logger.debug(f"Identified caption indices: {caption_indices}")
        return caption_indices

    def _apply_spacing_docx(
        self,
        input_path: str,
        output_path: str,
        spacing_options: Dict[str, bool]
    ) -> Dict[str, Any]:
        """
        Apply spacing to DOCX document parts.

        Args:
            input_path: Input file path
            output_path: Output file path
            spacing_options: Dictionary with boolean flags

        Returns:
            Result information
        """
        logger.info("Applying spacing to DOCX document")

        try:
            doc = Document(input_path)
            spacing_applied = 0

            # Build the structural model once for this operation
            structure = self._build_structure(doc, source_path=input_path)

            # Apply spacing based on options
            if spacing_options.get('paragraphs', False):
                logging.debug(f"Identifying PARAGRAPHS for spacing application")
                paragraphs = self._identify_paragraphs(doc, structure)
                for idx in paragraphs:
                    logger.debug(f"Applying spacing to paragraph {idx}: '{doc.paragraphs[idx].text[:30]}...'")
                    self._add_paragraph_spacing(doc.paragraphs[idx])
                    spacing_applied += 1
                logger.info(f"Applied spacing to {len(paragraphs)} paragraphs")

            if spacing_options.get('sentences', False):
                logger.debug(f"Identifying SENTENCES for spacing application")
                sentences_to_process = self._identify_sentences(doc, structure)
                logger.debug(f"Sentences identified for spacing: {sentences_to_process[:3]} (showing first 3)")
                # Process in reversed order to avoid index issues when modifying paragraphs
                for para_idx, sentences in reversed(sentences_to_process):
                    logger.debug(f"Applying sentence spacing to paragraph {para_idx} with sentences: {sentences[:3]} (showing first 3)")
                    self._add_sentence_spacing(doc.paragraphs[para_idx])
                    spacing_applied += 1
                logger.info(f"Applied spacing to sentences in {len(sentences_to_process)} paragraphs")

            # Save document
            doc.save(output_path)
            logger.info(f"DOCX spacing completed: {spacing_applied} spacings applied")

            return {
                'success': True,
                'output_path': output_path,
                'format': 'docx',
                'spacing_applied': spacing_applied
            }

        except Exception as e:
            logger.error(f"DOCX spacing error: {str(e)}")
            raise FormattingException(f"DOCX spacing failed: {str(e)}")

    def _apply_framing_docx(
        self,
        input_path: str,
        output_path: str,
        framing_options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Apply framing (borders) to DOCX document parts using table encapsulation.

        This method encapsulates paragraphs or sections in 1x1 tables with customizable borders.
        It preserves all content including:
        - Text formatting (bold, italic, underline, colors)
        - Hyperlinks and references
        - Images and embedded objects
        - Paragraph styles and alignment
        - Spacing before/after paragraphs

        Args:
            input_path: Input file path
            output_path: Output file path
            framing_options: Dictionary with framing options:
                - sections: bool - Frame entire sections
                - paragraphs: bool - Frame individual paragraphs
                - subparagraphs: bool - Frame subparagraphs
                - sentences: bool - Frame sentences
                - use_tables: bool - Use table encapsulation (default: True)
                - border_width: int - Border width in eighths of a point (default: 8)
                - border_color: str - Border color in hex without # (default: "000000")
                - border_style: str - Border style: single, double, dashed, etc. (default: "single")
                - cell_margin: int - Cell margin in twips (default: 100)
                - preserve_spacing: bool - Preserve paragraph spacing (default: True)
                - filter: dict - Filtering options:
                    - style_names: List[str] - Only frame these styles
                    - exclude_headings: bool - Exclude headings (default: True)
                    - exclude_empty: bool - Exclude empty paragraphs (default: True)
                    - min_length: int - Minimum text length (default: 0)
                    - has_marker: str - Only frame paragraphs containing this text

        Returns:
            Result information including number of frames applied

        Raises:
            FormattingException: If framing fails
        """
        logger.info("Applying framing to DOCX document")

        try:
            doc = Document(input_path)
            borders_applied = 0

            # Build the structural model once for this operation
            structure = self._build_structure(doc, source_path=input_path)

            # Extract options with defaults
            use_tables = framing_options.get('use_tables', True)
            border_width = framing_options.get('border_width', None)
            border_color = framing_options.get('border_color', None)
            border_style = framing_options.get('border_style', None)
            cell_margin = framing_options.get('cell_margin', None)
            preserve_spacing = framing_options.get('preserve_spacing', True)
            filter_options = framing_options.get('filter', None)

            logger.info(f"Framing mode: {'Table encapsulation' if use_tables else 'Paragraph borders'}")
            if use_tables:
                logger.info(f"Border settings - Width: {border_width or self.DEFAULT_TABLE_BORDER_WIDTH}, "
                           f"Color: {border_color or self.DEFAULT_BORDER_COLOR}, "
                           f"Style: {border_style or self.DEFAULT_BORDER_STYLE}")

            # Resolve border colour once. Style/width are resolved per part so
            # sections, paragraphs and sentences are visually distinguishable,
            # unless the request overrides them explicitly.
            border_color = (border_color or self.DEFAULT_BORDER_COLOR).lstrip('#')

            def _borders_for(part_style: str, part_width: int) -> Tuple[str, int]:
                return (border_style or part_style, border_width or part_width)

            if not use_tables:
                borders_applied += self._apply_framing_paragraph_borders(
                    doc, structure, framing_options, filter_options
                )
            else:
                borders_applied += self._apply_framing_tables(
                    doc, structure, framing_options, filter_options,
                    border_color, cell_margin, preserve_spacing, _borders_for
                )

            # Save document
            doc.save(output_path)

            logger.info(f"DOCX framing completed: {borders_applied} frames applied")

            return {
                'success': True,
                'output_path': output_path,
                'format': 'docx',
                'borders_applied': borders_applied
            }

        except Exception as e:
            logger.error(f"DOCX framing error: {str(e)}")
            raise FormattingException(f"DOCX framing failed: {str(e)}")

    def _apply_framing_paragraph_borders(
        self,
        doc: Document,
        structure: DocumentStructure,
        framing_options: Dict[str, Any],
        filter_options: Optional[Dict[str, Any]],
    ) -> int:
        """Legacy fallback: draw borders directly on paragraphs (no tables)."""
        borders_applied = 0
        indices = set()

        if framing_options.get('sections', False):
            for start_idx, end_idx, _ in self._identify_sections(doc, structure):
                indices.update(range(start_idx, end_idx + 1))
        if framing_options.get('paragraphs', False):
            indices.update(self._identify_paragraphs(doc, structure))
        if framing_options.get('subparagraphs', False):
            indices.update(idx for idx, _ in self._identify_subparagraphs(doc, structure))
        if framing_options.get('sentences', False):
            indices.update(idx for idx, _ in self._identify_sentences(doc, structure))

        for idx in sorted(indices):
            para = doc.paragraphs[idx]
            if self._should_frame_paragraph(para, filter_options):
                self._add_paragraph_border(para)
                borders_applied += 1
        return borders_applied

    def _apply_framing_tables(
        self,
        doc: Document,
        structure: DocumentStructure,
        framing_options: Dict[str, Any],
        filter_options: Optional[Dict[str, Any]],
        border_color: str,
        cell_margin: Optional[int],
        preserve_spacing: bool,
        borders_for,
    ) -> int:
        """
        Encapsulate document parts in single-cell tables at the right granularity.

        Each part is framed with its own table unit:
        - ``sections``  -> one table wrapping all paragraphs of the section,
        - ``paragraphs``/``subparagraphs`` -> one table per paragraph,
        - ``sentences`` -> one table per sentence.

        Every table is followed by an empty paragraph so adjacent tables are not
        merged by Word and the blocks stay visually separated.

        Precedence when several options are enabled: sections > paragraphs >
        subparagraphs > sentences. A paragraph already consumed by a
        higher-priority part is not framed again.
        """
        # Collect operations as (kind, payload) using live paragraph objects so
        # that tree mutations during execution do not invalidate integer indices.
        operations: List[Tuple[str, Any]] = []
        consumed = set()  # ids of paragraph elements already assigned to a part

        def _take(para) -> bool:
            key = id(para._element)
            if key in consumed:
                return False
            if not self._should_frame_paragraph(para, filter_options):
                return False
            consumed.add(key)
            return True

        if framing_options.get('sections', False):
            sections = self._identify_sections(doc, structure)
            logger.info(f"Identified {len(sections)} sections to frame")
            for start_idx, end_idx, _ in sections:
                paras = []
                for idx in range(start_idx, end_idx + 1):
                    para = doc.paragraphs[idx]
                    if _take(para):
                        paras.append(para)
                if paras:
                    operations.append(('section', paras))

        if framing_options.get('paragraphs', False):
            paragraphs = self._identify_paragraphs(doc, structure)
            logger.info(f"Identified {len(paragraphs)} paragraphs to frame")
            for idx in paragraphs:
                para = doc.paragraphs[idx]
                if _take(para):
                    operations.append(('paragraph', [para]))

        if framing_options.get('subparagraphs', False):
            subparagraphs = self._identify_subparagraphs(doc, structure)
            logger.info(f"Identified {len(subparagraphs)} subparagraphs to frame")
            for para_idx, _ in subparagraphs:
                para = doc.paragraphs[para_idx]
                if _take(para):
                    operations.append(('paragraph', [para]))

        if framing_options.get('sentences', False):
            sentence_map = self._identify_sentences(doc, structure)
            logger.info(f"Identified sentences in {len(sentence_map)} paragraphs to frame")
            for para_idx, sentences in sentence_map:
                if not sentences:
                    continue
                para = doc.paragraphs[para_idx]
                if _take(para):
                    operations.append(('sentences', (para, sentences)))

        # Execute bottom-to-top so inserting/removing elements never shifts the
        # positions of parts we have not processed yet.
        body = list(doc.element.body)

        def _first_para_element(op):
            _kind, payload = op
            para = payload[0] if _kind != 'sentences' else payload[0]
            return para._element

        def _pos(op):
            el = _first_para_element(op)
            try:
                return body.index(el)
            except ValueError:
                return -1

        operations.sort(key=_pos, reverse=True)

        borders_applied = 0
        for kind, payload in operations:
            try:
                if kind == 'section':
                    style, width = borders_for(self.SECTION_BORDER_STYLE, self.SECTION_BORDER_WIDTH)
                    self._encapsulate_paragraphs_in_table(
                        payload, doc, border_width=width, border_color=border_color,
                        border_style=style, cell_margin=cell_margin,
                        preserve_spacing=preserve_spacing,
                    )
                elif kind == 'paragraph':
                    style, width = borders_for(self.PARAGRAPH_BORDER_STYLE, self.PARAGRAPH_BORDER_WIDTH)
                    self._encapsulate_paragraphs_in_table(
                        payload, doc, border_width=width, border_color=border_color,
                        border_style=style, cell_margin=cell_margin,
                        preserve_spacing=preserve_spacing,
                    )
                elif kind == 'sentences':
                    para, sentences = payload
                    style, width = borders_for(self.SENTENCE_BORDER_STYLE, self.SENTENCE_BORDER_WIDTH)
                    self._encapsulate_sentences_in_tables(
                        para, sentences, doc, border_width=width, border_color=border_color,
                        border_style=style, cell_margin=cell_margin,
                        preserve_spacing=preserve_spacing,
                    )
                borders_applied += 1
            except Exception as e:
                logger.error(f"Failed to frame {kind}: {str(e)}")
                # Continue with the remaining parts

        logger.info(f"Encapsulated {len(operations)} parts in tables")
        return borders_applied

    def _apply_framing_pdf(
        self,
        input_path: str,
        output_path: str,
        framing_options: Dict[str, bool]
    ) -> Dict[str, Any]:
        """
        Apply framing (borders) to PDF document parts.

        TODO: This is a placeholder for future PDF border implementation.
        PDF border application requires more complex text extraction and positioning.

        Args:
            input_path: Input file path
            output_path: Output file path
            framing_options: Dictionary with boolean flags

        Returns:
            Result information
        """
        # logger.info("Applying framing to PDF document (basic implementation)")
        # logger.warning("PDF framing support is limited and under development")
        #
        # try:
        #     # Open PDF
        #     doc = pymupdf.open(input_path)
        #     output_doc = pymupdf.open()
        #
        #     borders_applied = 0
        #
        #     for page in doc:
        #         # Create new page
        #         new_page = output_doc.new_page(width=page.rect.width, height=page.rect.height)
        #
        #         # Copy original content
        #         new_page.show_pdf_page(new_page.rect, doc, page.number)
        #
        #         # TODO: Implement text block identification and border drawing
        #         # This requires:
        #         # 1. Extract text blocks with positions
        #         # 2. Identify sections/paragraphs/sentences based on positions and content
        #         # 3. Draw rectangles around identified blocks
        #
        #         # Placeholder: Draw border around entire page as example
        #         if any(framing_options.values()):
        #             rect = pymupdf.Rect(50, 50, page.rect.width - 50, page.rect.height - 50)
        #             new_page.draw_rect(rect, color=(0, 0, 0), width=0.5)
        #             borders_applied += 1
        #
        #     # Save output
        #     output_doc.save(output_path)
        #     output_doc.close()
        #     doc.close()
        #
        #     logger.info(f"PDF framing completed: {borders_applied} borders applied (basic)")
        #
        #     return {
        #         'success': True,
        #         'output_path': output_path,
        #         'format': 'pdf',
        #         'borders_applied': borders_applied,
        #         'framing_options': framing_options,
        #         'note': 'PDF framing is in development - currently applies page-level borders'
        #     }
        #
        # except Exception as e:
        #     logger.error(f"PDF framing error: {str(e)}")
        #     raise FormattingException(f"PDF framing failed: {str(e)}")
        pass

    def _apply_keywords_docx(
        self,
        input_path: str,
        output_path: str,
        keyword_options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Estrae parole chiave dalle sezioni di un documento DOCX e le inserisce come paragrafi formattati.

        Questa funzione analizza un documento DOCX per identificare le sezioni strutturali
        (tramite _identify_sections) e per ciascuna estrae le parole chiave più rilevanti
        utilizzando Ollama come metodo principale, con fallback automatico a spaCy in caso
        di indisponibilità del servizio o errori.

        Flusso di esecuzione:
        1. Caricamento del documento DOCX
        2. Identificazione delle sezioni tramite _identify_sections() che rileva:
           - Pattern di testo (es. " A.", " B.")
           - Stili Heading (Heading 2, 3, etc.)
           - Formattazione speciale (grassetto + font > 11pt)
        3. Per ogni sezione identificata:
           a. Estrazione del testo combinato di tutti i paragrafi della sezione
           b. Tentativo di estrazione keywords con Ollama (metodo principale)
           c. In caso di fallimento, utilizzo di spaCy come fallback
           d. Formattazione delle keywords: "Parole chiave: keyword1, keyword2, ..."
           e. Inserimento come nuovo paragrafo formattato dopo il titolo della sezione
        4. Salvataggio del documento modificato
        5. Restituzione metadati completi sull'operazione

        Formattazione delle keywords inserite:
        - Stile: Italico
        - Dimensione font: 10pt (20 half-points in OpenXML)
        - Posizione: Subito dopo il titolo della sezione
        - Formato testo: "Parole chiave: parola1, parola2, parola3"

        Gestione errori e fallback:
        - Se Ollama non è disponibile: usa spaCy per tutte le sezioni
        - Se Ollama fallisce per una sezione specifica: usa spaCy per quella sezione
        - Se nessuna sezione è identificata: salva documento invariato e restituisce warning
        - Ogni errore viene loggato con dettagli per debugging

        Args:
            input_path (str): Percorso assoluto del file DOCX di input
            output_path (str): Percorso assoluto dove salvare il file DOCX processato
            keyword_options (Dict[str, Any]): Dizionario con opzioni di estrazione:
                - max_keywords (int): Numero di parole chiave per sezione (range: 1-10).
                                     Default: 5
                - include_proper_nouns (bool): Include nomi propri nell'estrazione
                                               (utilizzato solo dal fallback spaCy).
                                               Default: True
                - model (str, optional): Nome del modello Ollama specifico da utilizzare
                                        (es. 'llama2', 'mistral', 'phi').
                                        Se omesso, usa il modello di default configurato.
                                        Default: None

        Returns:
            Dict[str, Any]: Dizionario con informazioni dettagliate sul risultato:
                {
                    'success': bool,                    # True se operazione completata
                    'output_path': str,                 # Path del file generato
                    'format': str,                      # Formato documento ('docx')
                    'sections_processed': int,          # Numero di sezioni processate
                    'total_keywords': int,              # Totale keywords estratte
                    'keyword_options': Dict[str, Any],  # Opzioni utilizzate
                    'extraction_method': str,           # Metodo usato: 'Ollama', 'spaCy',
                                                       # o 'Ollama with spaCy fallback'
                    'ollama_used': bool,               # True se Ollama è stato usato
                    'spacy_fallback_used': bool        # True se fallback è stato attivato
                }

                In caso di nessuna sezione trovata, include anche:
                    'note': str                        # Messaggio descrittivo

        Raises:
            FormattingException: Se si verifica un errore durante:
                - Caricamento del documento
                - Elaborazione delle sezioni
                - Salvataggio del file
                - Qualsiasi altro errore di processing
        """
        logger.info("Applying keyword extraction to DOCX document")

        try:
            from app.services.keyword_service import get_keyword_service
            from app.services.ollama_service import get_ollama_service

            # Load document
            doc = Document(input_path)

            # Build the structural model once for this operation
            structure = self._build_structure(doc, source_path=input_path)

            # Get options
            max_keywords = keyword_options.get('max_keywords', 5)
            include_proper_nouns = keyword_options.get('include_proper_nouns', True)
            ollama_model = keyword_options.get('model', None)

            sections_processed = 0
            total_keywords_extracted = 0
            ollama_used = False
            spacy_fallback_used = False

            # Identify sections using the advanced method
            sections = self._identify_sections(doc, structure)
            logger.info(f"Identified {len(sections)} sections for keyword extraction")

            if not sections:
                logger.warning("No sections identified in document")
                doc.save(output_path)
                return {
                    'success': True,
                    'output_path': output_path,
                    'format': 'docx',
                    'sections_processed': 0,
                    'total_keywords': 0,
                    'keyword_options': keyword_options,
                    'note': 'No sections found in document'
                }

            # Try to get Ollama service
            ollama_service = None
            try:
                ollama_service = get_ollama_service()
                use_ollama = True
                logger.info("Ollama service available, will use for keyword extraction")
            except Exception as e:
                use_ollama = False
                logger.warning(f"Ollama service not available: {str(e)}. Will use spaCy fallback")

            # Get spaCy service as fallback
            keyword_service = get_keyword_service()

            logger.debug("Starting section processing loop")
            logger.debug("#################################")
            logger.debug(f"Sections to process: {len(sections)}")
            logger.debug(f"Sections identified: {[(s[0], s[1]) for s in sections]}")
            logger.debug("#################################")

            # IMPORTANTE: Process sections in REVERSE order to avoid index shift problems
            # When we insert a paragraph before a section, all subsequent paragraph indices shift by +1
            # By processing from the last section to the first, we ensure that:
            # - Inserting keywords in section N doesn't affect the indices of sections 0..N-1
            # - The indices remain valid throughout the entire processing loop
            logger.info(f"Processing {len(sections)} sections in REVERSE order to preserve paragraph indices")

            # Process each section (in reverse order)
            for section_num, (start_idx, end_idx, section_text) in enumerate(reversed(sections), 1):
                actual_section_num = len(sections) - section_num + 1  # For logging
                logger.debug(f"Processing section {actual_section_num}/{len(sections)}: paragraphs {start_idx}-{end_idx}")

                # Get the first paragraph of the section to use as title for logging
                first_para = doc.paragraphs[start_idx].text
                logger.debug(f"Section starts with: '{first_para[:50]}...'")

                if not section_text.strip():
                    logger.debug(f"Section {start_idx}-{end_idx} has no text, skipping")
                    continue

                # Extract keywords using Ollama or spaCy
                keywords = []

                if use_ollama:
                    try:
                        keywords = ollama_service.extract_keywords(
                            text=section_text,
                            max_keywords=max_keywords,
                            model=ollama_model,
                            use_cache=True
                        )
                        ollama_used = True
                        logger.debug(f"Ollama extracted keywords: {keywords}")
                    except Exception as e:
                        logger.warning(f"Ollama keyword extraction failed: {str(e)}. Falling back to spaCy")
                        use_ollama = False  # Disable for remaining sections

                # Fallback to spaCy if Ollama failed or not available
                if not keywords:
                    keywords = keyword_service.extract_keywords(
                        text=section_text,
                        max_keywords=max_keywords,
                        include_proper_nouns=include_proper_nouns
                    )
                    spacy_fallback_used = True
                    logger.debug(f"spaCy extracted keywords: {keywords}")

                if keywords:
                    # Detect the section language so the prefix is localised
                    # (e.g. "Parole chiave" for Italian, "Keywords" for English).
                    section_language = keyword_service.detect_language(section_text)

                    # Format keywords using keyword_service
                    # This returns: "<localised label>: keyword1, keyword2, ..."
                    keyword_text = keyword_service.format_keywords(keywords, language=section_language)
                    logger.debug(f"Formatted keywords for section '{first_para[:30]}...': {keyword_text}")

                    # Insert keyword paragraph right after the section start
                    start_para = doc.paragraphs[start_idx]

                    # Create new paragraph element
                    new_para_element = OxmlElement('w:p')

                    # Add paragraph properties with spacing after, so the
                    # keywords are visually separated from the section content.
                    pPr = OxmlElement('w:pPr')
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:after'), str(self.KEYWORD_SPACING_AFTER_TWIPS))
                    pPr.append(spacing)
                    new_para_element.append(pPr)

                    # Create run element with formatting properties
                    run_element = OxmlElement('w:r')

                    # Apply formatting: italic and 10pt font
                    rPr = OxmlElement('w:rPr')
                    italic = OxmlElement('w:i')
                    sz = OxmlElement('w:sz')
                    sz.set(qn('w:val'), '20')  # 10pt = 20 half-points
                    rPr.append(italic)
                    rPr.append(sz)
                    run_element.append(rPr)

                    # Create text element with keyword text
                    text_element = OxmlElement('w:t')
                    text_element.text = keyword_text
                    run_element.append(text_element)

                    # Add run to paragraph
                    new_para_element.append(run_element)

                    # Insert the new paragraph before the section title
                    start_para._element.addprevious(new_para_element)

                    sections_processed += 1
                    total_keywords_extracted += len(keywords)

                    logger.info(f"Added keywords before section '{first_para[:50]}...': {keyword_text}")

            # Save document
            doc.save(output_path)

            # Prepare result metadata
            extraction_method = "Ollama" if ollama_used else "spaCy"
            if ollama_used and spacy_fallback_used:
                extraction_method = "Ollama with spaCy fallback"

            logger.info(f"DOCX keyword extraction completed: {sections_processed} sections processed, "
                       f"{total_keywords_extracted} keywords extracted using {extraction_method}")

            return {
                'success': True,
                'output_path': output_path,
                'format': 'docx',
                'sections_processed': sections_processed,
                'total_keywords': total_keywords_extracted,
                'keyword_options': keyword_options,
                'extraction_method': extraction_method,
                'ollama_used': ollama_used,
                'spacy_fallback_used': spacy_fallback_used
            }

        except Exception as e:
            logger.error(f"DOCX keyword extraction error: {str(e)}")
            raise FormattingException(f"DOCX keyword extraction failed: {str(e)}")

    def _apply_highlighting_docx(
        self,
        input_path: str,
        output_path: str,
        highlighting_options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Apply part-of-speech text formatting to DOCX document.

        This method processes each paragraph in the document, analyzes the text using spaCy
        to identify parts of speech, and applies formatting (color, font, style) based on
        user-specified options.

        Args:
            input_path: Path to input DOCX file
            output_path: Path to save output DOCX file
            highlighting_options: Dictionary containing:
                - enabled: bool - Whether formatting is enabled
                - color: str - Text color in hex format (e.g., "#FF0000")
                - style: str - Text styles: 'bold', 'italic', 'underline', or combinations like 'bold,italic'
                - font_size: int - Font size in points (6-72)
                - font_family: str - Font family name (e.g., 'Times New Roman', 'Arial')
                - nouns: bool - Whether to format nouns
                - verbs: bool - Whether to format verbs
                - adjectives: bool - Whether to format adjectives
                - adverbs: bool - Whether to format adverbs

        Returns:
            Dictionary with result information:
                - success: bool - Whether operation succeeded
                - output_path: str - Path to output file
                - format: str - File format ("docx")
                - words_formatted: int - Total words formatted
                - paragraphs_processed: int - Number of paragraphs processed
                - pos_stats: dict - Statistics per part of speech
                - highlighting_options: dict - Options used

        Raises:
            FormattingException: If document processing fails
        """
        logger.info("Applying part-of-speech text formatting to DOCX document")

        try:
            from app.services.keyword_service import get_keyword_service

            # Load document
            doc = Document(input_path)

            # Build the structural model once for this operation
            structure = self._build_structure(doc, source_path=input_path)

            # Get keyword service for POS analysis
            keyword_service = get_keyword_service()

            # Identify sections (this excludes headings)
            sections = self._identify_sections(doc, structure)

            if not sections:
                logger.warning("No sections found in document. Nothing to format.")
                return {
                    'success': True,
                    'output_path': output_path,
                    'format': 'docx',
                    'words_formatted': 0,
                    'paragraphs_processed': 0,
                    'pos_stats': {'nouns': 0, 'verbs': 0, 'adjectives': 0, 'adverbs': 0},
                    'highlighting_options': highlighting_options
                }

            # Extract options
            color = highlighting_options.get('color', self.DEFAULT_HIGHLIGHT_COLOR).lstrip('#')
            style_str = highlighting_options.get('style', None)
            font_size = highlighting_options.get('font_size', None)
            font_family = highlighting_options.get('font_family', None)
            highlight_nouns = highlighting_options.get('nouns', False)
            highlight_verbs = highlighting_options.get('verbs', False)
            highlight_adjectives = highlighting_options.get('adjectives', False)
            highlight_adverbs = highlighting_options.get('adverbs', False)

            # Parse styles
            apply_bold = False
            apply_italic = False
            apply_underline = False
            if style_str:
                styles = [s.strip().lower() for s in style_str.split(',')]
                apply_bold = 'bold' in styles
                apply_italic = 'italic' in styles
                apply_underline = 'underline' in styles

            # Convert hex color to RGB
            try:
                r = int(color[0:2], 16)
                g = int(color[2:4], 16)
                b = int(color[4:6], 16)
            except ValueError:
                logger.warning(f"Invalid color format: {color}, using default black")
                r, g, b = self.DEFAULT_TEXT_COLOR_RGB

            # Statistics
            words_formatted = 0
            paragraphs_processed = 0
            pos_stats = {
                'nouns': 0,
                'verbs': 0,
                'adjectives': 0,
                'adverbs': 0
            }

            logger.info(f"Processing document with color RGB({r}, {g}, {b})")
            logger.info(f"Font settings - Family: {font_family}, Size: {font_size}, Styles: bold={apply_bold}, italic={apply_italic}, underline={apply_underline}")
            logger.info(f"POS to format - Nouns: {highlight_nouns}, Verbs: {highlight_verbs}, "
                       f"Adjectives: {highlight_adjectives}, Adverbs: {highlight_adverbs}")
            logger.info(f"Found {len(sections)} sections to process")

            # Build a set of paragraph indices that are part of sections (not headings)
            section_paragraph_indices = set()
            for start_idx, end_idx, _ in sections:
                for i in range(start_idx, end_idx + 1):
                    section_paragraph_indices.add(i)

            logger.info(f"Processing {len(section_paragraph_indices)} paragraphs (excluding headings)")

            # Process only paragraphs that are part of sections
            for para_idx, paragraph in enumerate(doc.paragraphs):
                # Skip if this paragraph is not part of a section (it's likely a heading)
                if para_idx not in section_paragraph_indices:
                    logger.debug(f"Skipping paragraph {para_idx} (heading or non-section content)")
                    continue

                # IMPORTANT: Also skip if this paragraph IS a heading (title)
                # Even if it's in section_paragraph_indices, headings should not be processed
                if structure.is_heading(para_idx):
                    logger.debug(f"Skipping paragraph {para_idx} (is a heading)")
                    continue

                # Skip paragraphs starting with "Parole chiave"
                if paragraph.text.strip().startswith("Parole chiave"):
                    logger.debug(f"Skipping paragraph {para_idx} (starts with 'Parole chiave')")
                    continue

                if not paragraph.text.strip():
                    continue

                # CRITICAL: Skip paragraphs that contain complex fields (citations/references)
                # Processing paragraphs with fields can cause them to be moved or corrupted
                paragraph_has_complex_field = False
                for run in paragraph.runs:
                    if run._element is not None:
                        for child in run._element:
                            tag_name = child.tag
                            if 'fldChar' in tag_name or 'instrText' in tag_name or 'fldData' in tag_name:
                                paragraph_has_complex_field = True
                                break
                        if paragraph_has_complex_field:
                            break

                if paragraph_has_complex_field:
                    logger.debug(f"Skipping paragraph {para_idx} (contains citation/reference fields)")
                    continue

                paragraphs_processed += 1

                # Analyze text with spaCy
                tokens = keyword_service.analyze_pos(paragraph.text)

                if not tokens:
                    continue

                # Store paragraph-level formatting
                original_style = paragraph.style
                original_alignment = paragraph.alignment


                # Build a map of character positions to format decisions
                # Map directly from tokens with POS decisions
                char_format_decisions = {}

                for token in tokens:
                    # Skip punctuation and spaces
                    if token['is_punct'] or token['is_space']:
                        continue

                    token_start = token.get('start_char', 0)
                    token_end = token.get('end_char', 0)
                    pos = token['pos']

                    # Determine if this token should be formatted based on its POS
                    should_format = False
                    if highlight_nouns and pos in ['NOUN', 'PROPN']:
                        should_format = True
                        pos_stats['nouns'] += 1
                    elif highlight_verbs and pos == 'VERB':
                        should_format = True
                        pos_stats['verbs'] += 1
                    elif highlight_adjectives and pos == 'ADJ':
                        should_format = True
                        pos_stats['adjectives'] += 1
                    elif highlight_adverbs and pos == 'ADV':
                        should_format = True
                        pos_stats['adverbs'] += 1

                    # Mark all characters in this token with the formatting decision
                    if should_format:
                        words_formatted += 1
                        for char_pos in range(token_start, token_end):
                            char_format_decisions[char_pos] = True

                # Build a list of runs with metadata about what to do with each
                runs_info = []
                current_char_pos = 0

                for run in paragraph.runs:
                    run_text = run.text
                    run_length = len(run_text)

                    # Check if this run contains complex field elements (citations, references)
                    has_complex_field = False
                    if run._element is not None:
                        for child in run._element:
                            tag_name = child.tag
                            if 'fldChar' in tag_name or 'instrText' in tag_name or 'fldData' in tag_name:
                                has_complex_field = True
                                break

                    # Store run info
                    runs_info.append({
                        'run': run,
                        'element': run._element,  # Keep reference to original element
                        'start_pos': current_char_pos,
                        'end_pos': current_char_pos + run_length,
                        'text': run_text,
                        'has_complex_field': has_complex_field,
                        'original_formatting': {
                            'font_name': run.font.name,
                            'font_size': run.font.size,
                            'font_color': run.font.color.rgb if run.font.color.rgb else None,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline
                        }
                    })

                    current_char_pos += run_length

                # Get paragraph element for manipulating XML
                paragraph_element = paragraph._element

                # First, remove all existing runs from the paragraph
                for run in paragraph.runs:
                    run._element.getparent().remove(run._element)

                # Now rebuild runs in the correct order
                for run_info in runs_info:
                    # If this run has complex fields, reinsert the original element
                    if run_info['has_complex_field']:
                        logger.debug(f"Preserving run with complex field (citation/reference)")
                        paragraph_element.append(run_info['element'])
                        continue

                    run_start = run_info['start_pos']
                    run_text = run_info['text']
                    orig_fmt = run_info['original_formatting']

                    if not run_text:
                        # Empty run, skip it
                        continue

                    # Split this run into segments based on char_format_decisions
                    char_idx = 0
                    while char_idx < len(run_text):
                        char_pos = run_start + char_idx
                        should_format_segment = char_format_decisions.get(char_pos, False)

                        # Find the end of this segment (consecutive chars with same decision)
                        segment_start = char_idx
                        while (char_idx < len(run_text) and
                               char_format_decisions.get(run_start + char_idx, False) == should_format_segment):
                            char_idx += 1

                        # Create a new run element for this segment
                        segment_text = run_text[segment_start:char_idx]

                        # Use paragraph.add_run which adds to the end (which maintains order since we process sequentially)
                        new_run = paragraph.add_run(segment_text)

                        # Apply original formatting
                        if orig_fmt['font_name']:
                            new_run.font.name = orig_fmt['font_name']
                        if orig_fmt['font_size']:
                            new_run.font.size = orig_fmt['font_size']
                        if orig_fmt['font_color']:
                            new_run.font.color.rgb = orig_fmt['font_color']
                        if orig_fmt['bold'] is not None:
                            new_run.bold = orig_fmt['bold']
                        if orig_fmt['italic'] is not None:
                            new_run.italic = orig_fmt['italic']
                        if orig_fmt['underline'] is not None:
                            new_run.underline = orig_fmt['underline']

                        # Apply new formatting if this segment should be formatted
                        if should_format_segment:
                            new_run.font.color.rgb = RGBColor(r, g, b)

                            if font_family:
                                new_run.font.name = font_family

                            if font_size:
                                new_run.font.size = Pt(font_size)

                            if apply_bold:
                                new_run.bold = True
                            if apply_italic:
                                new_run.italic = True
                            if apply_underline:
                                new_run.underline = True


                # Restore paragraph formatting
                paragraph.style = original_style
                if original_alignment:
                    paragraph.alignment = original_alignment

            # Save document
            doc.save(output_path)

            logger.info(f"DOCX text formatting completed: {paragraphs_processed} paragraphs processed, "
                       f"{words_formatted} words formatted")
            logger.info(f"POS statistics: {pos_stats}")

            return {
                'success': True,
                'output_path': output_path,
                'format': 'docx',
                'words_formatted': words_formatted,
                'paragraphs_processed': paragraphs_processed,
                'pos_stats': pos_stats,
                'highlighting_options': highlighting_options
            }

        except Exception as e:
            logger.error(f"DOCX text formatting error: {str(e)}")
            raise FormattingException(f"DOCX text formatting failed: {str(e)}")

    def get_available_styles(self, file_format: str) -> Dict[str, Any]:
        """
        Get available formatting options for a file format.

        Args:
            file_format: File format (docx, pdf, txt)

        Returns:
            Dictionary with available options
        """
        if file_format not in self.SUPPORTED_FORMATS:
            return {
                'format': file_format,
                'supported': False,
                'options': []
            }

        return {
            'format': file_format,
            'supported': True,
            'options': self.SUPPORTED_FORMATS[file_format],
            'option_details': {
                'font_name': {
                    'type': 'string',
                    'description': 'Font family name',
                    'examples': ['Arial', 'Times New Roman', 'Calibri']
                },
                'font_size': {
                    'type': 'number',
                    'description': 'Font size in points',
                    'examples': [10, 11, 12, 14, 16]
                },
                'font_color': {
                    'type': 'string',
                    'description': 'Font color in hex format',
                    'examples': ['#000000', '#FF0000', '#0000FF']
                },
                'bold': {
                    'type': 'boolean',
                    'description': 'Apply bold formatting'
                },
                'italic': {
                    'type': 'boolean',
                    'description': 'Apply italic formatting'
                },
                'alignment': {
                    'type': 'string',
                    'description': 'Text alignment',
                    'examples': ['left', 'center', 'right', 'justify']
                }
            }
        }

    def _copy_run_formatting(self, source_run, new_run) -> None:
        """Copy character formatting (and hyperlinks) from one run to another."""
        if source_run.bold is not None:
            new_run.bold = source_run.bold
        if source_run.italic is not None:
            new_run.italic = source_run.italic
        if source_run.underline is not None:
            new_run.underline = source_run.underline
        if source_run.font.name:
            new_run.font.name = source_run.font.name
        if source_run.font.size:
            new_run.font.size = source_run.font.size
        if source_run.font.color and source_run.font.color.rgb:
            new_run.font.color.rgb = source_run.font.color.rgb

        # Copy hyperlinks and other complex child elements at the XML level.
        if source_run._element is not None:
            for child in source_run._element:
                if 'hyperlink' in str(child.tag).lower():
                    new_run._element.append(child)

    def _copy_paragraph_into(self, dest_paragraph, source_paragraph, preserve_spacing: bool = True) -> None:
        """Copy style, alignment, spacing and all runs from source into dest."""
        try:
            dest_paragraph.style = source_paragraph.style
        except Exception:
            logger.debug(f"Could not copy paragraph style: {source_paragraph.style}")

        if source_paragraph.alignment:
            dest_paragraph.alignment = source_paragraph.alignment

        if preserve_spacing:
            src_fmt = source_paragraph.paragraph_format
            if src_fmt.space_before is not None:
                dest_paragraph.paragraph_format.space_before = src_fmt.space_before
            if src_fmt.space_after is not None:
                dest_paragraph.paragraph_format.space_after = src_fmt.space_after

        for run in source_paragraph.runs:
            new_run = dest_paragraph.add_run(run.text)
            self._copy_run_formatting(run, new_run)

    def _insert_table_at(self, doc, anchor_element, before: bool = True):
        """Create a 1x1 table and insert it relative to ``anchor_element``.

        Returns the ``(table, cell)`` pair. The table is inserted immediately
        before (default) or after the anchor element in the document tree.
        """
        table = doc.add_table(rows=1, cols=1)
        table_element = table._element
        if before:
            anchor_element.addprevious(table_element)
        else:
            anchor_element.addnext(table_element)
        return table, table.rows[0].cells[0]

    @staticmethod
    def _add_empty_paragraph_after(element) -> None:
        """Insert an empty paragraph right after ``element`` for separation.

        The trailing blank line keeps consecutive tables from being merged by
        Word into a single grid and gives the framed blocks visual breathing
        room.
        """
        element.addnext(OxmlElement('w:p'))

    def _encapsulate_paragraphs_in_table(
        self,
        paragraphs: List,
        doc,
        border_width: Optional[int] = None,
        border_color: Optional[str] = None,
        border_style: Optional[str] = None,
        cell_margin: Optional[int] = None,
        preserve_spacing: bool = True
    ) -> None:
        """
        Encapsulate one or more consecutive paragraphs in a single 1x1 table.

        All provided paragraphs are moved into the same table cell (each keeps
        its own line/paragraph inside the cell), the table is inserted at the
        position of the first paragraph, the originals are removed and an empty
        paragraph is appended after the table for visual separation.

        Args:
            paragraphs: List of python-docx paragraphs to wrap in one cell.
            doc: python-docx document object.
            border_width: Border width in eighths of a point.
            border_color: Border color in hex format without #.
            border_style: Border style ('single', 'double', 'dashed', ...).
            cell_margin: Cell margin in twips.
            preserve_spacing: Whether to preserve paragraph spacing.
        """
        if not paragraphs:
            return

        border_width = border_width if border_width is not None else self.DEFAULT_TABLE_BORDER_WIDTH
        border_color = border_color if border_color is not None else self.DEFAULT_BORDER_COLOR
        border_style = border_style if border_style is not None else self.DEFAULT_BORDER_STYLE
        cell_margin = cell_margin if cell_margin is not None else self.DEFAULT_TABLE_CELL_MARGIN

        first_element = paragraphs[0]._element

        # Insert the table just before the first paragraph.
        table, cell = self._insert_table_at(doc, first_element, before=True)
        table_element = table._element

        # Copy each source paragraph into the single cell.
        for i, src in enumerate(paragraphs):
            dest = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            self._copy_paragraph_into(dest, src, preserve_spacing)

        # Apply borders and margins.
        self._set_table_borders(table, border_width, border_color, border_style)
        self._set_table_cell_margins(table, cell_margin)

        # Remove the original paragraphs from the body.
        for src in paragraphs:
            el = src._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

        # Trailing empty paragraph so tables don't merge and blocks stay separated.
        self._add_empty_paragraph_after(table_element)

        logger.debug(
            f"Encapsulated {len(paragraphs)} paragraph(s) in a table: "
            f"'{cell.paragraphs[0].text[:50]}...'"
        )

    def _encapsulate_paragraph_in_table(
        self,
        paragraph,
        doc,
        border_width: Optional[int] = None,
        border_color: Optional[str] = None,
        border_style: Optional[str] = None,
        cell_margin: Optional[int] = None,
        preserve_spacing: bool = True
    ) -> None:
        """Encapsulate a single paragraph in a 1x1 table (see the list variant)."""
        self._encapsulate_paragraphs_in_table(
            [paragraph], doc,
            border_width=border_width,
            border_color=border_color,
            border_style=border_style,
            cell_margin=cell_margin,
            preserve_spacing=preserve_spacing,
        )

    def _encapsulate_sentences_in_tables(
        self,
        paragraph,
        sentences: List[str],
        doc,
        border_width: Optional[int] = None,
        border_color: Optional[str] = None,
        border_style: Optional[str] = None,
        cell_margin: Optional[int] = None,
        preserve_spacing: bool = True
    ) -> int:
        """
        Replace a paragraph with one 1x1 table per sentence.

        Each sentence is placed in its own single-cell table, followed by an
        empty paragraph so the sentence boxes stay visually separated and are
        not merged by Word. Sentence text is re-emitted as plain runs inheriting
        the paragraph style/alignment; per-run character formatting inside a
        sentence is not preserved (the paragraph is split by text).

        Returns:
            The number of sentence tables created.
        """
        if not sentences:
            return 0

        border_width = border_width if border_width is not None else self.DEFAULT_TABLE_BORDER_WIDTH
        border_color = border_color if border_color is not None else self.DEFAULT_BORDER_COLOR
        border_style = border_style if border_style is not None else self.DEFAULT_BORDER_STYLE
        cell_margin = cell_margin if cell_margin is not None else self.DEFAULT_TABLE_CELL_MARGIN

        original_style = paragraph.style
        original_alignment = paragraph.alignment
        para_element = paragraph._element

        created = 0
        for sentence in sentences:
            text = sentence.strip()
            if not text:
                continue

            # Insert the sentence table before the (soon removed) paragraph so
            # the resulting order matches the original reading order.
            table, cell = self._insert_table_at(doc, para_element, before=True)
            dest = cell.paragraphs[0]
            try:
                dest.style = original_style
            except Exception:
                logger.debug(f"Could not copy paragraph style: {original_style}")
            if original_alignment:
                dest.alignment = original_alignment
            dest.add_run(text)

            self._set_table_borders(table, border_width, border_color, border_style)
            self._set_table_cell_margins(table, cell_margin)
            self._add_empty_paragraph_after(table._element)
            created += 1

        # Remove the original paragraph now that its sentences are framed.
        parent = para_element.getparent()
        if parent is not None:
            parent.remove(para_element)

        logger.debug(f"Encapsulated {created} sentence(s) in individual tables")
        return created

    def _set_table_borders(
        self,
        table,
        border_width: int,
        border_color: str,
        border_style: str
    ) -> None:
        """
        Set borders for a table.

        Args:
            table: python-docx table object
            border_width: Border width in eighths of a point
            border_color: Border color in hex format without #
            border_style: Border style (single, double, dashed, etc.)
        """
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Remove existing borders if any
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)

        # Create new borders
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), border_style)
            border.set(qn('w:sz'), str(border_width))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def _set_table_cell_margins(self, table, margin: int) -> None:
        """
        Set cell margins for a table.

        Args:
            table: python-docx table object
            margin: Margin in twips (1/1440 inch)
        """
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Remove existing margins if any
        existing_margins = tblPr.find(qn('w:tblCellMar'))
        if existing_margins is not None:
            tblPr.remove(existing_margins)

        # Create cell margins
        tblCellMar = OxmlElement('w:tblCellMar')

        for side in ['top', 'left', 'bottom', 'right']:
            margin_elem = OxmlElement(f'w:{side}')
            margin_elem.set(qn('w:w'), str(margin))
            margin_elem.set(qn('w:type'), 'dxa')  # dxa = twentieths of a point
            tblCellMar.append(margin_elem)

        tblPr.append(tblCellMar)

    def _should_frame_paragraph(
        self,
        paragraph,
        filter_options: Optional[Dict[str, Any]] = None
    ) -> bool:
        """
        Determine if a paragraph should be framed based on filter options.

        Args:
            paragraph: python-docx paragraph object
            filter_options: Dictionary with filtering criteria:
                - style_names: List of style names to include (e.g., ['Normal', 'Body Text'])
                - exclude_headings: Whether to exclude heading paragraphs
                - exclude_empty: Whether to exclude empty paragraphs
                - min_length: Minimum text length to include
                - has_marker: Text marker that paragraph must contain

        Returns:
            bool: True if paragraph should be framed, False otherwise
        """
        if filter_options is None:
            return True

        # Check if empty and should exclude
        if filter_options.get('exclude_empty', True) and not paragraph.text.strip():
            return False

        # Check minimum length
        min_length = filter_options.get('min_length', 0)
        if len(paragraph.text) < min_length:
            return False

        # Check if heading and should exclude
        if filter_options.get('exclude_headings', True):
            if self._is_heading(paragraph):
                return False

        # Check style names
        style_names = filter_options.get('style_names', None)
        if style_names:
            if paragraph.style.name not in style_names:
                return False

        # Check for required marker
        has_marker = filter_options.get('has_marker', None)
        if has_marker:
            if has_marker not in paragraph.text:
                return False

        return True


# Singleton instance
_formatting_service_instance = None


def get_formatting_service() -> FormattingService:
    """
    Get singleton instance of FormattingService.

    Returns:
        FormattingService: Singleton instance
    """
    global _formatting_service_instance
    if _formatting_service_instance is None:
        _formatting_service_instance = FormattingService()
    return _formatting_service_instance

