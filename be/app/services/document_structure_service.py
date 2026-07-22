"""
Document Structure Service - Canonical structural model for DOCX documents.

This module centralises the logic that identifies the structural role of each
paragraph in a DOCX document (document title, section heading, content
paragraph, caption, ...).

The design is inspired by the way structure is made *explicit* in Markdown
(``#`` = document title, ``##`` = section, blank lines separate paragraphs).
In a DOCX the very same information already exists natively through the
paragraph *style outline level* (``Title`` / ``Heading 1`` / ``Heading 2`` ...),
so we classify blocks by style instead of guessing from font size or bold runs.

Classification strategy:
    * Option A (primary): read the paragraph style outline level.
    * Option B (fallback): when a document carries no heading styles at all,
      convert DOCX -> Markdown with ``mammoth`` and recover the heading
      structure from the ``#`` markers, aligning them back to the DOCX
      paragraphs by text. If ``mammoth`` is unavailable, degrade to a
      font-size / bold heuristic.

Every :class:`Block` keeps its originating ``index`` in ``doc.paragraphs`` so
that downstream formatting operations can still write back to the exact DOCX
paragraph.
"""
import logging
import re
from dataclasses import dataclass, field
from enum import Enum
from typing import List, Optional, Tuple

from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class BlockType(Enum):
    """Structural role of a paragraph."""
    DOCUMENT_TITLE = "document_title"    # Word 'Title' style (the document main title)
    SECTION_HEADING = "section_heading"  # Heading 1 / Markdown '#'
    PARAGRAPH_TITLE = "paragraph_title"  # Heading 2 and deeper / Markdown '##', '###', ...
    PARAGRAPH = "paragraph"              # Regular content paragraph
    CAPTION = "caption"                  # Image / figure caption
    EMPTY = "empty"                      # Blank paragraph


@dataclass
class Block:
    """A single classified paragraph."""
    index: int                 # Index inside doc.paragraphs
    type: BlockType
    level: Optional[int]       # 0 = document title, >=1 = heading depth, None otherwise
    text: str
    style_name: str

    @property
    def is_heading(self) -> bool:
        """True for any block that acts as a section delimiter."""
        return self.type in (
            BlockType.DOCUMENT_TITLE,
            BlockType.SECTION_HEADING,
            BlockType.PARAGRAPH_TITLE,
        )


@dataclass
class Section:
    """A run of content paragraphs delimited by headings."""
    start_index: int           # First content paragraph of the section
    end_index: int             # Last content paragraph of the section
    text: str
    heading_index: Optional[int] = None   # Heading that opened the section (if any)
    heading_level: Optional[int] = None
    body_indices: List[int] = field(default_factory=list)

    def as_tuple(self) -> Tuple[int, int, str]:
        """Backward-compatible representation used by the formatting service."""
        return self.start_index, self.end_index, self.text


class DocumentStructure:
    """Immutable structural view over a python-docx ``Document``."""

    def __init__(self, blocks: List[Block], strategy: str) -> None:
        self.blocks = blocks
        self.strategy = strategy  # 'style' (Option A), 'markdown' or 'heuristic' (Option B)
        self._by_index = {b.index: b for b in blocks}

    # -- per-paragraph queries -------------------------------------------------

    def block_at(self, index: int) -> Optional[Block]:
        return self._by_index.get(index)

    def is_heading(self, index: int) -> bool:
        block = self._by_index.get(index)
        return bool(block and block.is_heading)

    def is_main_title(self, index: int) -> bool:
        block = self._by_index.get(index)
        return bool(block and block.type == BlockType.DOCUMENT_TITLE)

    def is_section_heading(self, index: int) -> bool:
        block = self._by_index.get(index)
        return bool(block and block.type == BlockType.SECTION_HEADING)

    def is_paragraph_title(self, index: int) -> bool:
        block = self._by_index.get(index)
        return bool(block and block.type == BlockType.PARAGRAPH_TITLE)

    # -- collection queries ----------------------------------------------------

    def main_title_indices(self) -> List[int]:
        return [b.index for b in self.blocks if b.type == BlockType.DOCUMENT_TITLE]

    def section_heading_indices(self) -> List[int]:
        return [b.index for b in self.blocks if b.type == BlockType.SECTION_HEADING]

    def paragraph_title_indices(self) -> List[int]:
        return [b.index for b in self.blocks if b.type == BlockType.PARAGRAPH_TITLE]

    def caption_indices(self) -> List[int]:
        return [b.index for b in self.blocks if b.type == BlockType.CAPTION]

    def content_paragraph_indices(self) -> List[int]:
        """
        All content paragraphs of the document, i.e. every non-empty paragraph
        that is not a heading, a document title or a caption.

        Content is captured regardless of the heading level it sits under (or
        even before any heading), so the introductory part of a document and
        the paragraphs under top-level headings are grouped just like those
        under lower-level section headings.
        """
        return [
            block.index
            for block in self.blocks
            if block.type == BlockType.PARAGRAPH and block.text.strip()
        ]

    def sections(self) -> List[Section]:
        """
        A section is the block of content between two headings (or from a
        heading to the end of the document). Headings themselves are not part
        of the section body.
        """
        sections: List[Section] = []
        current_start: Optional[int] = None
        current_heading: Optional[Block] = None
        body_texts: List[str] = []
        body_indices: List[int] = []

        for block in self.blocks:
            if block.type == BlockType.EMPTY:
                if current_start is not None:
                    body_texts.append(block.text)
                continue

            if block.is_heading:
                if current_start is not None and body_indices:
                    sections.append(Section(
                        start_index=current_start,
                        end_index=block.index - 1,
                        text='\n'.join(body_texts),
                        heading_index=current_heading.index if current_heading else None,
                        heading_level=current_heading.level if current_heading else None,
                        body_indices=list(body_indices),
                    ))
                current_start = None
                current_heading = block
                body_texts = []
                body_indices = []
            else:
                if current_start is None:
                    current_start = block.index
                body_texts.append(block.text)
                body_indices.append(block.index)

        if current_start is not None and body_indices:
            sections.append(Section(
                start_index=current_start,
                end_index=self.blocks[-1].index if self.blocks else current_start,
                text='\n'.join(body_texts),
                heading_index=current_heading.index if current_heading else None,
                heading_level=current_heading.level if current_heading else None,
                body_indices=list(body_indices),
            ))

        return sections


class DocumentStructureService:
    """Builds a :class:`DocumentStructure` from a python-docx ``Document``."""

    # Style-name matching. Built-in styles are exposed by python-docx with their
    # fixed English names ("Heading 1", "Title"); the extra localisations below
    # only matter for custom styles named in another language.
    _HEADING_RE = re.compile(
        r'(?:heading|titolo|titre|titel|t[íi]tulo|[uü]berschrift|encabezado)\s*(\d+)',
        re.IGNORECASE,
    )

    # Document-title style names across common Word localisations
    # (English "Title", Italian "Titolo", French "Titre", German "Titel",
    # Spanish/Portuguese "Título"/"Titulo"). A trailing digit excludes localised
    # heading styles such as "Titolo 1", which are handled by ``_HEADING_RE``.
    _TITLE_STYLE_RE = re.compile(
        r'^(?:title|titolo|titre|titel|t[íi]tulo)(?!\s*\d)\b',
        re.IGNORECASE,
    )

    # Caption style names across common Word localisations
    # (English, Italian, French, German, Spanish, Portuguese).
    _CAPTION_STYLE_RE = re.compile(
        r'caption|didascal|legend|l[eé]gende|beschriftung|leyenda|epígrafe|epigrafe',
        re.IGNORECASE,
    )

    # Caption text pattern, e.g. "Figura 2: ...", "Figure 2. ...", "Tabella 1 - ...".
    _CAPTION_TEXT_RE = re.compile(
        r'^\s*(figure|figura|fig\.?|table|tabella|tab\.?|immagine|foto|grafico|'
        r'grafik|schema|illustrazione|diagramma|abbildung|tabla|imagen)\s*\.?\s*'
        r'\d+\s*[:.\)\-\u2013\u2014]',
        re.IGNORECASE,
    )

    # Heuristic fallback thresholds (used only when no heading styles exist)
    HEADING_FONT_SIZE_THRESHOLD = 14.0

    def __init__(self) -> None:
        logger.info("Document structure service initialized")

    # -- public API ------------------------------------------------------------

    def build(self, doc, source_path: Optional[str] = None) -> DocumentStructure:
        """
        Classify every paragraph of ``doc`` and return its structure.

        Args:
            doc: python-docx ``Document`` object.
            source_path: Path to the original ``.docx`` on disk. Required to
                enable the Markdown (mammoth) fallback; when omitted the
                fallback degrades to the font-size/bold heuristic.
        """
        paragraphs = doc.paragraphs

        # Option A: rely on native heading styles when present.
        if self._has_style_headings(paragraphs):
            blocks = [self._classify_by_style(i, p) for i, p in enumerate(paragraphs)]
            logger.info("Document structure built using style outline levels (Option A)")
            return DocumentStructure(blocks, strategy='style')

        # Option B: no heading styles -> recover structure from Markdown.
        logger.info("No heading styles found; falling back to Markdown/heuristic detection (Option B)")
        return self._build_fallback(paragraphs, source_path)

    # -- single-paragraph classification --------------------------------------

    def classify_paragraph(self, para, use_heuristic: bool = True) -> Block:
        """
        Classify a single paragraph in isolation (index unknown).

        The style outline level is checked first. When the style carries no
        heading information and ``use_heuristic`` is True, a font-size/bold
        heuristic is used as a last resort.
        """
        text = para.text
        style_name = para.style.name
        if not text.strip():
            return Block(-1, BlockType.EMPTY, None, text, style_name)

        role, level = self._heading_role(para)
        if role is not None:
            return Block(-1, role, level, text, style_name)
        if self._is_caption_text(text):
            return Block(-1, BlockType.CAPTION, None, text, style_name)

        if use_heuristic:
            return self._classify_heuristic(-1, para)
        return Block(-1, BlockType.PARAGRAPH, None, text, style_name)

    # -- Option A --------------------------------------------------------------

    _HEADING_TYPES = (
        BlockType.DOCUMENT_TITLE,
        BlockType.SECTION_HEADING,
        BlockType.PARAGRAPH_TITLE,
    )

    def _has_style_headings(self, paragraphs) -> bool:
        for para in paragraphs:
            _type, _level = self._heading_role(para)
            if _type in self._HEADING_TYPES:
                return True
        return False

    @staticmethod
    def _role_for_depth(depth: int) -> BlockType:
        """Map a 0-based heading depth to its structural role.

        depth 0 -> section heading (Heading 1 / Markdown '#'),
        depth >= 1 -> paragraph title (Heading 2+ / Markdown '##'+).

        The document title is never derived from a heading depth: it comes
        exclusively from the dedicated ``Title`` style (see ``_style_role``).
        """
        if depth <= 0:
            return BlockType.SECTION_HEADING
        return BlockType.PARAGRAPH_TITLE

    @staticmethod
    def _outline_level(para) -> Optional[int]:
        """Return the 0-based Word outline level of ``para``, or None.

        Word records a heading's depth in ``w:pPr/w:outlineLvl`` (0 = Heading 1,
        1 = Heading 2, ...) independently of the style *name*, which may be
        localised or renamed. The value is read from the paragraph's own
        properties first, then from the style definition following the
        ``basedOn`` inheritance chain.
        """
        outline_qn = qn('w:outlineLvl')
        val_qn = qn('w:val')

        p_pr = getattr(para, '_p', None)
        p_pr = p_pr.pPr if p_pr is not None else None
        if p_pr is not None:
            node = p_pr.find(outline_qn)
            if node is not None and node.get(val_qn) is not None:
                return int(node.get(val_qn))

        style = para.style
        visited = set()
        while style is not None and id(style) not in visited:
            visited.add(id(style))
            element = getattr(style, 'element', None)
            if element is not None:
                s_pr = element.find(qn('w:pPr'))
                if s_pr is not None:
                    node = s_pr.find(outline_qn)
                    if node is not None and node.get(val_qn) is not None:
                        return int(node.get(val_qn))
            style = getattr(style, 'base_style', None)
        return None

    def _style_role(self, style_name: str) -> Tuple[Optional[BlockType], Optional[int]]:
        """Map a style *name* to a (BlockType, level) role, or (None, None)."""
        low = (style_name or '').strip().lower()
        # Heading number is checked first so localised heading styles such as
        # "Titolo 1" are not mistaken for the document "Title" style.
        match = self._HEADING_RE.search(low)
        if match:
            level = int(match.group(1))
            return self._role_for_depth(level - 1), level
        if self._TITLE_STYLE_RE.search(low):
            return BlockType.DOCUMENT_TITLE, 0
        if self._CAPTION_STYLE_RE.search(low):
            return BlockType.CAPTION, None
        return None, None

    def _heading_role(self, para) -> Tuple[Optional[BlockType], Optional[int]]:
        """Resolve the heading role of ``para`` from its style name or outline level.

        The style name is matched first (``Title`` / ``Heading N``). When the
        style carries no recognisable name -- e.g. a custom or localised title
        style -- the Word outline level is used as a fallback so that any style
        actually configured as a heading is still classified correctly.
        """
        role, level = self._style_role(para.style.name)
        if role in self._HEADING_TYPES:
            return role, level

        depth = self._outline_level(para)
        if depth is not None:
            return self._role_for_depth(depth), depth + 1

        return role, level

    def _is_caption_text(self, text: str) -> bool:
        """Detect captions from their textual pattern (e.g. 'Figura 2: ...')."""
        return bool(text) and bool(self._CAPTION_TEXT_RE.match(text))

    def _classify_by_style(self, index: int, para) -> Block:
        style_name = para.style.name
        text = para.text
        if not text.strip():
            return Block(index, BlockType.EMPTY, None, text, style_name)

        role, level = self._heading_role(para)
        if role is not None:
            return Block(index, role, level, text, style_name)
        if self._is_caption_text(text):
            return Block(index, BlockType.CAPTION, None, text, style_name)
        return Block(index, BlockType.PARAGRAPH, None, text, style_name)

    # -- Option B (fallback) ---------------------------------------------------

    def _build_fallback(self, paragraphs, source_path: Optional[str]) -> DocumentStructure:
        heading_levels = self._markdown_heading_levels(source_path)

        if heading_levels:
            blocks = self._classify_with_heading_map(paragraphs, heading_levels)
            logger.info("Document structure built from Markdown conversion (Option B)")
            return DocumentStructure(blocks, strategy='markdown')

        blocks = [self._classify_heuristic(i, p) for i, p in enumerate(paragraphs)]
        logger.info("Document structure built using font-size/bold heuristic (Option B)")
        return DocumentStructure(blocks, strategy='heuristic')

    def _markdown_heading_levels(self, source_path: Optional[str]) -> dict:
        """
        Convert the DOCX to Markdown with mammoth and return a mapping
        ``{normalized_text: markdown_heading_level}`` for every ``#`` heading.
        Returns an empty dict when the path/mammoth are unavailable or fail.
        """
        if not source_path:
            logger.warning("No source path provided; skipping Markdown fallback")
            return {}

        try:
            import mammoth
        except ImportError:
            logger.warning("mammoth not installed; skipping Markdown fallback")
            return {}

        try:
            with open(source_path, 'rb') as handle:
                result = mammoth.convert_to_markdown(handle)
            markdown = result.value
        except Exception as exc:  # pragma: no cover - defensive
            logger.warning(f"Markdown conversion failed: {exc}")
            return {}

        heading_re = re.compile(r'^(#{1,6})\s+(.*)$')
        levels: dict = {}
        for line in markdown.splitlines():
            match = heading_re.match(line.strip())
            if match:
                level = len(match.group(1))
                key = self._normalize(match.group(2))
                if key:
                    levels[key] = level
        return levels

    def _classify_with_heading_map(self, paragraphs, heading_levels: dict) -> List[Block]:
        blocks: List[Block] = []
        for index, para in enumerate(paragraphs):
            text = para.text
            style_name = para.style.name
            if not text.strip():
                blocks.append(Block(index, BlockType.EMPTY, None, text, style_name))
                continue

            level = heading_levels.get(self._normalize(text))
            if level is not None:
                blocks.append(Block(index, self._role_for_depth(level - 1), level, text, style_name))
            elif self._TITLE_STYLE_RE.search((style_name or '').lower()):
                blocks.append(Block(index, BlockType.DOCUMENT_TITLE, 0, text, style_name))
            elif self._CAPTION_STYLE_RE.search((style_name or '').lower()) or self._is_caption_text(text):
                blocks.append(Block(index, BlockType.CAPTION, None, text, style_name))
            else:
                blocks.append(Block(index, BlockType.PARAGRAPH, None, text, style_name))
        return blocks

    def _classify_heuristic(self, index: int, para) -> Block:
        text = para.text
        style_name = para.style.name
        if not text.strip():
            return Block(index, BlockType.EMPTY, None, text, style_name)

        if self._CAPTION_STYLE_RE.search((style_name or '').lower()) or self._is_caption_text(text):
            return Block(index, BlockType.CAPTION, None, text, style_name)

        font_size = self._paragraph_font_size(para)
        is_bold = bool(para.runs) and all(run.font.bold for run in para.runs)

        if font_size is not None and font_size >= self.HEADING_FONT_SIZE_THRESHOLD:
            # The largest fonts are treated as the document title, smaller
            # heading-sized fonts as section headings (Heading 1 equivalent).
            block_type = BlockType.DOCUMENT_TITLE if font_size >= self.HEADING_FONT_SIZE_THRESHOLD + 4 else BlockType.SECTION_HEADING
            level = 0 if block_type == BlockType.DOCUMENT_TITLE else 1
            return Block(index, block_type, level, text, style_name)
        if is_bold:
            # Bold-only emphasis is treated as a lower-level paragraph title.
            return Block(index, BlockType.PARAGRAPH_TITLE, 2, text, style_name)

        return Block(index, BlockType.PARAGRAPH, None, text, style_name)

    # -- helpers ---------------------------------------------------------------

    @staticmethod
    def _paragraph_font_size(para) -> Optional[float]:
        try:
            if para.style.font.size:
                return para.style.font.size.pt
        except (AttributeError, TypeError):
            pass
        for run in para.runs:
            try:
                if run.font.size:
                    return run.font.size.pt
            except (AttributeError, TypeError):
                continue
        return None

    @staticmethod
    def _normalize(text: str) -> str:
        """Normalise text for matching Markdown headings to DOCX paragraphs."""
        text = re.sub(r'[\\*_`#\[\]()]', '', text)   # strip markdown syntax/escapes
        text = re.sub(r'\s+', ' ', text)
        return text.strip().lower()


_document_structure_service_instance: Optional[DocumentStructureService] = None


def get_document_structure_service() -> DocumentStructureService:
    """Return the DocumentStructureService singleton."""
    global _document_structure_service_instance
    if _document_structure_service_instance is None:
        _document_structure_service_instance = DocumentStructureService()
        logger.info("Document structure service singleton instance created")
    return _document_structure_service_instance
